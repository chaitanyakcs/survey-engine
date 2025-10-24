"""Document parsing service for converting DOCX files to JSON using LLM."""

import json
import logging
from typing import Dict, Any, Optional, List
from io import BytesIO
from docx import Document
import replicate
from pydantic import ValidationError
import uuid
import time

from ..models.survey import SurveyCreate, SurveyCreateWithSections, Question, QuestionType, SurveySection
from ..config.settings import settings
from ..utils.error_messages import UserFriendlyError, get_api_configuration_error
from ..utils.llm_audit_decorator import LLMAuditContext
from ..services.llm_audit_service import LLMAuditService
from ..utils.json_generation_utils import parse_llm_json_response, get_json_optimized_hyperparameters, create_json_system_prompt, get_rfq_parsing_schema

logger = logging.getLogger(__name__)

class DocumentParsingError(Exception):
    """Exception raised when document parsing fails."""
    pass

class DocumentParser:
    """Service for parsing DOCX documents and converting to survey JSON."""

    def __init__(self, db_session: Optional[Any] = None, rfq_parsing_manager: Optional[Any] = None):
        if not settings.replicate_api_token:
            error_info = get_api_configuration_error()
            raise UserFriendlyError(
                message=error_info["message"],
                technical_details="REPLICATE_API_TOKEN environment variable is not set",
                action_required="Configure AI service provider (Replicate or OpenAI)"
            )
        self.replicate_client = replicate.Client(api_token=settings.replicate_api_token)
        self.model = self._get_rfq_parsing_model(db_session)  # Get model from RFQ parsing settings
        logger.info(f"🚀 [DocumentParser] Initialized with model: {self.model}")
        logger.info(f"🚀 [DocumentParser] Model source: DocumentParser.__init__")
        self.db_session = db_session
        self.rfq_parsing_manager = rfq_parsing_manager

    def _get_rfq_parsing_model(self, db_session: Optional[Any] = None) -> str:
        """Get RFQ parsing model from database settings with valid fallback"""
        # Valid Replicate model for RFQ parsing (never use settings.py)
        VALID_DEFAULT_MODEL = "openai/gpt-5"
        
        try:
            logger.info(f"🔍 [DocumentParser] Starting RFQ parsing model selection...")
            logger.info(f"🔍 [DocumentParser] Database session available: {bool(db_session)}")
            
            if db_session:
                from src.services.settings_service import SettingsService
                settings_service = SettingsService(db_session)
                rfq_settings = settings_service.get_rfq_parsing_settings()
                
                logger.info(f"🔍 [DocumentParser] Database RFQ parsing settings: {rfq_settings}")
                
                if rfq_settings and 'parsing_model' in rfq_settings:
                    model = rfq_settings['parsing_model']
                    logger.info(f"🔧 [DocumentParser] Using model from RFQ parsing settings: {model}")
                    logger.info(f"🔧 [DocumentParser] Model source: RFQ_PARSING_DATABASE")
                    return model
                else:
                    logger.warning(f"⚠️ [DocumentParser] No RFQ parsing settings found in database")
                    logger.info(f"🔧 [DocumentParser] Using valid default model: {VALID_DEFAULT_MODEL}")
                    logger.info(f"🔧 [DocumentParser] Model source: VALID_DEFAULT")
                    return VALID_DEFAULT_MODEL
            else:
                logger.warning(f"⚠️ [DocumentParser] No database session available")
                logger.info(f"🔧 [DocumentParser] Using valid default model: {VALID_DEFAULT_MODEL}")
                logger.info(f"🔧 [DocumentParser] Model source: VALID_DEFAULT_NO_DB")
                return VALID_DEFAULT_MODEL
        except Exception as e:
            logger.error(f"❌ [DocumentParser] Failed to get RFQ parsing model from database: {str(e)}")
            logger.info(f"🔧 [DocumentParser] Using valid default model: {VALID_DEFAULT_MODEL}")
            logger.info(f"🔧 [DocumentParser] Model source: VALID_DEFAULT_ERROR")
            return VALID_DEFAULT_MODEL

    async def _send_progress(self, session_id: str, stage: str, progress: int, message: str, details: Optional[str] = None, estimated_time: Optional[int] = None, content_preview: Optional[str] = None):
        """Send enhanced progress update via WebSocket if manager is available."""
        if self.rfq_parsing_manager and session_id:
            progress_data = {
                "type": "progress",
                "stage": stage,
                "progress": progress,
                "message": message,
                "details": details,
                "estimated_time": estimated_time,
                "content_preview": content_preview,
                "timestamp": time.time()
            }
            try:
                await self.rfq_parsing_manager.send_progress(session_id, progress_data)
                logger.debug(f"📤 [DocumentParser] Sent enhanced progress: {stage} ({progress}%) - {message}")
            except Exception as e:
                logger.warning(f"⚠️ [DocumentParser] Failed to send progress update: {str(e)}")
    
    async def extract_text_from_docx(self, docx_content: bytes, session_id: str = None) -> str:
        """Extract text content from DOCX file with detailed progress updates."""
        logger.info(f"📄 [Document Parser] Starting text extraction from DOCX, size: {len(docx_content)} bytes")
        try:
            doc = Document(BytesIO(docx_content))
            text_content = []
            
            # Send initial progress
            if session_id:
                await self._send_progress(session_id, "extracting", 5, 
                    "Opening document...", 
                    "Reading DOCX file structure", 
                    estimated_time=30)
            
            logger.info(f"📝 [Document Parser] Processing {len(doc.paragraphs)} paragraphs")
            paragraph_count = 0
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
                    paragraph_count += 1
                    if i < 5:  # Log first 5 paragraphs for debugging
                        logger.debug(f"📝 [Document Parser] Paragraph {i}: {paragraph.text.strip()[:100]}...")
                
                # Send progress updates every 10 paragraphs
                if session_id and i % 10 == 0 and i > 0:
                    progress = min(5 + (i / len(doc.paragraphs)) * 15, 20)
                    await self._send_progress(session_id, "extracting", int(progress),
                        f"Processing paragraphs... ({i}/{len(doc.paragraphs)})",
                        f"Found {paragraph_count} non-empty paragraphs",
                        estimated_time=25)
            
            logger.info(f"📊 [Document Parser] Processing {len(doc.tables)} tables")
            if session_id:
                await self._send_progress(session_id, "extracting", 20,
                    "Processing tables...",
                    f"Found {len(doc.tables)} tables to process",
                    estimated_time=20)
            
            # Also extract text from tables
            table_count = 0
            for table_idx, table in enumerate(doc.tables):
                logger.debug(f"📊 [Document Parser] Processing table {table_idx} with {len(table.rows)} rows")
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
                        table_count += 1
                
                # Send progress updates for tables
                if session_id and table_idx % 2 == 0:
                    progress = min(20 + (table_idx / len(doc.tables)) * 5, 25)
                    await self._send_progress(session_id, "extracting", int(progress),
                        f"Processing tables... ({table_idx + 1}/{len(doc.tables)})",
                        f"Extracted {table_count} table rows",
                        estimated_time=15)
            
            extracted_text = "\n".join(text_content)
            logger.info(f"✅ [Document Parser] Text extraction completed, total length: {len(extracted_text)} chars")
            
            # Send completion progress
            if session_id:
                word_count = len(extracted_text.split())
                await self._send_progress(session_id, "extracting", 25,
                    "Text extraction completed!",
                    f"Extracted {word_count} words from {paragraph_count} paragraphs and {table_count} table rows",
                    content_preview=extracted_text[:200] + "..." if len(extracted_text) > 200 else extracted_text,
                    estimated_time=10)
            
            return extracted_text
            
        except Exception as e:
            logger.error(f"❌ [Document Parser] Failed to extract text from DOCX: {str(e)}", exc_info=True)
            if session_id:
                await self._send_progress(session_id, "error", 0,
                    "Text extraction failed",
                    f"Error: {str(e)}",
                    estimated_time=0)
            raise DocumentParsingError(f"Failed to extract text from document: {str(e)}")
    
    def create_conversion_prompt(self, document_text: str, comments: List[Dict[str, Any]] = None) -> str:
        """Create the system prompt for LLM conversion with comment context."""
        
        # Define the expected JSON schema with sections format
        json_schema = {
            "title": "string (required)",
            "description": "string (required)", 
            "product_category": "string (electronics|appliances|healthcare_technology|enterprise_software|automotive|financial_services|hospitality)",
            "target_segment": "string (B2B decision makers|General consumers|Healthcare professionals|IT professionals|C-suite executives)",
            "research_goal": "string (pricing_research|feature_research|satisfaction_research|brand_research|market_sizing)",
            "methodologies": "array of strings (van_westendorp|gabor_granger|conjoint|maxdiff|brand_tracking|nps)",
            "estimated_time": "number (minutes)",
            "confidence_score": "number (0.0-1.0)",
            "sections": [
                {
                    "id": "integer (required, 1-7)",
                    "title": "string (required)",
                    "description": "string (optional)",
                    "introText": "object (optional) - for study introduction text",
                    "textBlocks": "array of objects (optional) - for additional text blocks",
                    "questions": [
                        {
                            "id": "string (required)",
                            "text": "string (required)",
                            "type": "string (REQUIRED: must be one of: multiple_choice|single_choice|text|scale|rating|yes_no|dropdown|matrix|ranking|numeric|date|boolean|file_upload|van_westendorp|gabor_granger|conjoint|maxdiff|instruction)",
                            "options": "array of strings (for choice questions)",
                            "scale_min": "number (for scale questions)",
                            "scale_max": "number (for scale questions)",
                            "scale_labels": "object with min/max keys (for scale questions, e.g. {'min': 'Poor', 'max': 'Excellent'})",
                            "required": "boolean",
                            "logic": "object (optional)",
                            "annotation": "object (optional) - include if there's a comment for this question",
                            "order": "integer (optional) - question order within section"
                        }
                    ]
                }
            ],
            "metadata": {
                "estimated_time": "number (minutes)",
                "target_responses": "number",
                "methodology_tags": "array of strings",
                "sections_count": "number"
            }
        }
        
        # Build comment context if available
        comment_context = ""
        if comments:
            comment_context = "\n\nCOMMENTS FROM DOCUMENT:\n"
            for i, comment in enumerate(comments):
                if comment.get('anchored_text'):
                    comment_context += f"{i+1}. Comment: \"{comment.get('text', '')}\"\n"
                    comment_context += f"   Anchored to: \"{comment.get('anchored_text', '')}\"\n"
                    if comment.get('author'):
                        comment_context += f"   Author: {comment.get('author')}\n"
                    comment_context += "\n"

        prompt = f"""You are an expert survey methodologist. Convert the following survey document to JSON format.

DOCUMENT TEXT:
{document_text}{comment_context}

INSTRUCTIONS:
1. Extract the survey structure and convert to the exact JSON schema provided below
2. CRITICAL: Question types MUST be one of: multiple_choice, single_choice, text, scale, rating, yes_no, dropdown, matrix, ranking, numeric, date, boolean, file_upload, van_westendorp, gabor_granger, conjoint, maxdiff, instruction
3. CRITICAL: If question type cannot be determined, use "instruction" instead of "unknown"
4. Detect methodologies: van_westendorp, gabor_granger, conjoint, maxdiff, brand_tracking, nps
5. Preserve question order and logic flow
6. If unsure about any field, use null rather than guessing
7. Generate a confidence_score (0.0-1.0) based on how well you could parse the document
8. CRITICAL: Always extract the COMPLETE question text from the document - never truncate based on comment anchors
9. CRITICAL: Check research objectives section for methodology requirements (e.g., "Most Important and Least Important features" indicates MaxDiff needed)
10. CRITICAL: Generate survey using SECTIONS format - this is MANDATORY:
   - Look for section headers like "SECTION A", "SECTION B", "Screener", "Introduction", etc.
   - Map sections to the 7-section structure: Sample Plan, Screener, Brand/Product Awareness, Concept Exposure, Methodology, Additional Questions, Programmer Instructions
   - If no explicit sections, organize questions logically into appropriate sections based on content
   - Each section must have: id (1-7), title, description, questions array
   - DO NOT use flat "questions" array - always use "sections" format
11. COMMENT MATCHING: If there are comments provided, match each comment to the question it was anchored to:
   - CRITICAL: Extract the COMPLETE question text from the document, not just the anchored portion
   - The "anchored_text" is only for matching comments to questions - it may be partial or truncated
   - Find the FULL question text in the document that contains or relates to the anchored text
   - If a comment matches a question, add an "annotation" field to that question with:
     - "comment": the comment text
     - "anchored_text": the exact text the comment was attached to (may be partial)
     - "author": the comment author (if available)
     - "date": the comment date (if available)
   - EXAMPLE: If anchored_text is "[EQ04] On a scale of 1 to 5 where 1 is 'Completely Unlikely' and 5 is 'Completely Likely', how likely would you be to f", 
     find the complete question like "How likely would you be to follow up and learn more about Product_B?" in the document

REQUIRED JSON SCHEMA:
{json.dumps(json_schema, indent=2)}

METHODOLOGY DETECTION HINTS:
- Van Westendorp PSM: Questions about "too cheap", "too expensive", "cheap", "expensive" price points
- Gabor-Granger: Sequential price acceptance questions
- Conjoint Analysis: Choice scenarios with multiple attributes
- MaxDiff: "Most important" and "Least important" selection tasks with features list
  CRITICAL: MaxDiff is ONE question per concept showing ALL features, not multiple questions
  CRITICAL: Extract actual feature names from concept description into 'features' array
  CRITICAL: Use 'features' field (not 'options') and set type to 'maxdiff'
  CRITICAL: Look for research objectives mentioning "Most Important and Least Important features"
- Brand Tracking: Brand awareness, consideration, preference questions
- NPS: "How likely are you to recommend" questions

QUESTION TEXT EXTRACTION RULES:
- CRITICAL: Use the ACTUAL question text from the document, not generic descriptions
- CRITICAL: Do not create generic descriptions like "Gabor-Granger: How likely would you be to purchase Product_A at varying price points?"
- CRITICAL: Extract the exact question text as written in the document
- EXAMPLE: If document has "CQ01a. On a scale of 1 to 5: where 1 is 'Definitely will not purchase' and 5 is 'Definitely will purchase', how likely would you be to purchase Product_A?", use that exact text
- EXAMPLE: If document has "Please highlight the MOST IMPORTANT and LEAST IMPORTANT features", use that exact text, not a generic description

UNKNOWN QUESTION TYPE HANDLING:
- CRITICAL: If question type cannot be determined, use "instruction" instead of "unknown"
- CRITICAL: Instructions are non-interactive text blocks that provide context or guidance
- CRITICAL: Instructions should not have options, scales, or interactive elements
- EXAMPLE: Text like "Please read the following information carefully" should be type "instruction"

SECTION MAPPING GUIDANCE:
- Section 1 (Sample Plan): Introduction text, study overview, recruitment criteria, quotas, confidentiality agreements
- Section 2 (Screener): SQ01-SQ09 type questions, demographics, qualification questions, participation checks
- Section 3 (Brand/Product Awareness): AQ00-AQ01 type questions, brand recall, usage patterns, product awareness
- Section 4 (Concept Exposure): BQ01-BQ07 type questions, concept introduction, impression ratings, product concepts
- Section 5 (Methodology): CQ01-CQ02 type questions, pricing studies, MaxDiff, conjoint, van_westendorp
- Section 6 (Additional Questions): DQ01-DQ05 type questions, demographics, supplementary questions, final questions
- Section 7 (Programmer Instructions): Technical notes, routing logic, implementation details, programming notes
- Look for patterns like "SQ01", "AQ00", "BQ01", "CQ01", "DQ01" to identify section boundaries
- If no clear section markers, group questions by topic: introduction → screening → awareness → concepts → methodology → demographics → technical

IMPORTANT: Return ONLY valid JSON that matches the schema exactly. No explanations or additional text.
"""
        return prompt
    
    def _post_process_question(self, question: Dict[str, Any]) -> Dict[str, Any]:
        """Clean up questions extracted from documents."""
        import copy
        import re
        
        original_question = copy.deepcopy(question)  # Preserve original
        q_type = question.get('type')
        
        try:
            if q_type == 'gabor_granger':
                return self._post_process_gabor_granger(question)
            elif q_type == 'maxdiff':
                return self._post_process_maxdiff(question)
            elif q_type == 'matrix' or q_type == 'matrix_likert':
                return self._post_process_matrix(question)
            elif q_type == 'instruction':
                return self._post_process_instruction(question)
            elif q_type == 'van_westendorp':
                return self._post_process_van_westendorp(question)
            elif q_type == 'dropdown':
                return self._post_process_dropdown(question)
            elif q_type == 'conjoint':
                return self._post_process_conjoint(question)
            elif q_type == 'yes_no':
                return self._post_process_yes_no(question)
            else:
                return self._clean_programming_instructions(question)
        except Exception as e:
            logger.error(f"Error post-processing {q_type} question {question.get('id')}: {e}", exc_info=True)
            return original_question  # Always return original on error

    def _clean_programming_instructions(self, question: Dict[str, Any]) -> Dict[str, Any]:
        """Remove common programming instructions from any question."""
        import re
        
        text = question.get('text', '')
        
        # Remove brackets with instructions
        text = re.sub(r'\[SHOW HYPERLINK[^\]]*\]', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\[HYPERLINK[^\]]*\]', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\[RANDOMIZE[^\]]*\]', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\[IF RESPONSE[^\]]*\]', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\[CAPTURE[^\]]*\]', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\[SELECT ONE[^\]]*\]', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\[ALLOW ONLY[^\]]*\]', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\[ANCHOR[^\]]*\]', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\[TERMINATE[^\]]*\]', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\[Thank & Terminate\]', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\[RECRUIT\]', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\[EXCLUSIVE\]', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\[COLLECT OE\]', '', text, flags=re.IGNORECASE)
        text = re.sub(r'QUOTAS:.*?(?=\n|$)', '', text, flags=re.IGNORECASE)
        
        # Clean up whitespace
        text = re.sub(r'\s+', ' ', text).strip()
        question['text'] = text
        
        return question

    def _post_process_gabor_granger(self, question: Dict[str, Any]) -> Dict[str, Any]:
        """Clean Gabor-Granger questions - extract prices, scale, product."""
        import copy
        import re
        
        original_question = copy.deepcopy(question)
        text = question.get('text', '')
        
        try:
            # Extract scale mapping
            scale_pattern = r'(?:scale of|from)\s*(\d+)\s*to\s*(\d+)\s*:?\s*where\s*(\d+)\s*is\s*[\'"]([^\'\"]+)[\'"].*?(?:and\s*)?(\d+)\s*is\s*[\'"]([^\'\"]+)[\'"]'
            scale_match = re.search(scale_pattern, text, re.IGNORECASE | re.DOTALL)
            
            if scale_match:
                _, _, low_num, low_label, high_num, high_label = scale_match.groups()
                question['scale_labels'] = {
                    low_num: low_label.strip(),
                    high_num: high_label.strip()
                }
            
            # Extract product name
            product_pattern = r'purchase\s+(Product_[A-Z]|GoPro[^?]+?)\??'
            product_match = re.search(product_pattern, text, re.IGNORECASE)
            product_name = product_match.group(1).strip() if product_match else 'this product'
            
            # Extract price points (pipe-delimited or space-separated)
            # Look for pattern like "Product_A at $249 | $299 | ..."
            price_section_pattern = rf'{re.escape(product_name)}\s+at\s+([\$\d\s\|,]+)'
            price_section_match = re.search(price_section_pattern, text, re.IGNORECASE)
            
            if price_section_match:
                price_text = price_section_match.group(1)
                prices = re.findall(r'\$[\d,]+(?:\.\d{2})?', price_text)
                # Remove duplicates while preserving order
                seen = set()
                unique_prices = []
                for price in prices:
                    if price not in seen:
                        seen.add(price)
                        unique_prices.append(price)
                question['options'] = unique_prices
            
            # Create clean question text
            clean_text = f"How likely would you be to purchase {product_name}?"
            
            # Add hyperlink context if present
            if 'click here to review' in text.lower():
                clean_text += " Click here to review the concept."
            
            question['text'] = clean_text
            
            # Safety checks
            if len(question.get('options', [])) == 0 and '$' in text:
                logger.warning(f"Price extraction failed for {question.get('id')}, keeping original")
                return original_question
                
            if len(question.get('text', '')) < 10:
                logger.warning(f"Cleaned text too short for {question.get('id')}, keeping original")
                return original_question
            
            logger.info(f"Successfully post-processed gabor_granger question {question.get('id')}")
            return question
            
        except Exception as e:
            logger.error(f"Error post-processing gabor_granger {question.get('id')}: {e}", exc_info=True)
            return original_question

    def _post_process_maxdiff(self, question: Dict[str, Any]) -> Dict[str, Any]:
        """Clean MaxDiff questions - extract features if highlighted."""
        import copy
        import re
        
        original_question = copy.deepcopy(question)
        text = question.get('text', '')
        
        try:
            # Remove programming instructions
            text = re.sub(r'\[SHOW CONCEPT[^\]]*\]', '', text, flags=re.IGNORECASE)
            text = re.sub(r'\[ALLOW RESPONDENTS[^\]]*\]', '', text, flags=re.IGNORECASE)
            
            # Extract product name
            product_pattern = r'(Product_[A-Z]|GoPro[^\s]+)'
            product_match = re.search(product_pattern, text)
            product_name = product_match.group(1) if product_match else 'the concept'
            
            # Create clean question
            clean_text = f"Please select the MOST IMPORTANT and LEAST IMPORTANT features of {product_name}."
            
            question['text'] = clean_text
            
            # Note: features array should be populated by generation service
            # from concept description, not from question text
            
            logger.info(f"Successfully post-processed maxdiff question {question.get('id')}")
            return question
            
        except Exception as e:
            logger.error(f"Error post-processing maxdiff {question.get('id')}: {e}", exc_info=True)
            return original_question

    def _post_process_matrix(self, question: Dict[str, Any]) -> Dict[str, Any]:
        """Clean matrix questions - remove routing/quota logic."""
        import copy
        import re
        
        original_question = copy.deepcopy(question)
        text = question.get('text', '')
        
        try:
            # Remove routing logic
            text = re.sub(r'\[IF [^\]]+\]', '', text, flags=re.IGNORECASE)
            text = re.sub(r'CLASSIFY AS[^\]]+', '', text, flags=re.IGNORECASE)
            text = re.sub(r'QUOTAS:[^$]+?(?=\[|$)', '', text, flags=re.IGNORECASE)
            
            # Clean up
            text = re.sub(r'\s+', ' ', text).strip()
            question['text'] = text
            
            result = self._clean_programming_instructions(question)
            logger.info(f"Successfully post-processed matrix question {question.get('id')}")
            return result
            
        except Exception as e:
            logger.error(f"Error post-processing matrix {question.get('id')}: {e}", exc_info=True)
            return original_question

    def _post_process_instruction(self, question: Dict[str, Any]) -> Dict[str, Any]:
        """Clean instruction questions - hide technical/termination screens."""
        import copy
        import re
        
        original_question = copy.deepcopy(question)
        text = question.get('text', '')
        
        try:
            # Flag technical instructions (not respondent-facing)
            technical_keywords = [
                'TERMINATING SCREEN',
                'RESPONDENT TYPE',
                'QUOTA',
                'CLASSIFICATION',
                'Thank & Terminate',
                'PROGRAMMING NOTE',
                'RANDOMIZATION:',
                'Implementation:'
            ]
            
            is_technical = any(keyword.lower() in text.lower() for keyword in technical_keywords)
            
            if is_technical:
                question['label'] = question.get('label', '') or 'Technical_Note'
                question['description'] = 'Programming instruction - not displayed to respondents'
            
            # Remove instruction markers
            text = re.sub(r'\[TERMINATING SCREEN\]', '', text, flags=re.IGNORECASE)
            text = re.sub(r'\[Show to [^\]]+\]', '', text, flags=re.IGNORECASE)
            
            question['text'] = text.strip()
            
            logger.info(f"Successfully post-processed instruction question {question.get('id')}")
            return question
            
        except Exception as e:
            logger.error(f"Error post-processing instruction {question.get('id')}: {e}", exc_info=True)
            return original_question

    def _post_process_van_westendorp(self, question: Dict[str, Any]) -> Dict[str, Any]:
        """Clean Van Westendorp questions - extract price sensitivity context."""
        import copy
        import re
        
        original_question = copy.deepcopy(question)
        text = question.get('text', '')
        
        try:
            # Remove programming instructions
            text = re.sub(r'\[SHOW HYPERLINK[^\]]*\]', '', text, flags=re.IGNORECASE)
            text = re.sub(r'\[RANDOMIZE[^\]]*\]', '', text, flags=re.IGNORECASE)
            text = re.sub(r'\[CAPTURE[^\]]*\]', '', text, flags=re.IGNORECASE)
            
            # Extract currency if present
            currency_match = re.search(r'[£$€¥₹]', text)
            currency = currency_match.group(0) if currency_match else '$'
            
            # Clean up text
            text = re.sub(r'\s+', ' ', text).strip()
            question['text'] = text
            
            # Add currency context
            if currency != '$':
                question['currency'] = currency
            
            logger.info(f"Successfully post-processed van_westendorp question {question.get('id')}")
            return question
            
        except Exception as e:
            logger.error(f"Error post-processing van_westendorp {question.get('id')}: {e}", exc_info=True)
            return original_question

    def _post_process_dropdown(self, question: Dict[str, Any]) -> Dict[str, Any]:
        """Clean dropdown questions - ensure options are properly formatted."""
        import copy
        import re
        
        original_question = copy.deepcopy(question)
        text = question.get('text', '')
        
        try:
            # Remove programming instructions
            text = re.sub(r'\[SHOW HYPERLINK[^\]]*\]', '', text, flags=re.IGNORECASE)
            text = re.sub(r'\[RANDOMIZE[^\]]*\]', '', text, flags=re.IGNORECASE)
            text = re.sub(r'\[SELECT ONE[^\]]*\]', '', text, flags=re.IGNORECASE)
            
            # Clean up text
            text = re.sub(r'\s+', ' ', text).strip()
            question['text'] = text
            
            # Ensure options are clean
            if 'options' in question and question['options']:
                cleaned_options = []
                for option in question['options']:
                    # Remove any programming markers from options
                    clean_option = re.sub(r'\[.*?\]', '', str(option)).strip()
                    if clean_option:
                        cleaned_options.append(clean_option)
                question['options'] = cleaned_options
            
            logger.info(f"Successfully post-processed dropdown question {question.get('id')}")
            return question
            
        except Exception as e:
            logger.error(f"Error post-processing dropdown {question.get('id')}: {e}", exc_info=True)
            return original_question

    def _post_process_conjoint(self, question: Dict[str, Any]) -> Dict[str, Any]:
        """Clean conjoint questions - extract attribute combinations."""
        import copy
        import re
        
        original_question = copy.deepcopy(question)
        text = question.get('text', '')
        
        try:
            # Remove programming instructions
            text = re.sub(r'\[SHOW CONCEPT[^\]]*\]', '', text, flags=re.IGNORECASE)
            text = re.sub(r'\[RANDOMIZE[^\]]*\]', '', text, flags=re.IGNORECASE)
            text = re.sub(r'\[ALLOW RESPONDENTS[^\]]*\]', '', text, flags=re.IGNORECASE)
            
            # Extract product name if present
            product_pattern = r'(Product_[A-Z]|GoPro[^\s]+)'
            product_match = re.search(product_pattern, text)
            product_name = product_match.group(1) if product_match else 'the product'
            
            # Create clean question
            clean_text = f"Please evaluate the following {product_name} combinations:"
            
            question['text'] = clean_text
            
            logger.info(f"Successfully post-processed conjoint question {question.get('id')}")
            return question
            
        except Exception as e:
            logger.error(f"Error post-processing conjoint {question.get('id')}: {e}", exc_info=True)
            return original_question

    def _post_process_yes_no(self, question: Dict[str, Any]) -> Dict[str, Any]:
        """Clean yes/no questions - remove programming instructions and ensure proper options."""
        import copy
        import re
        
        original_question = copy.deepcopy(question)
        text = question.get('text', '')
        
        try:
            # Remove programming instructions
            text = re.sub(r'\[FOR US ONLY\][^]]*', '', text, flags=re.IGNORECASE)
            text = re.sub(r'\[FOR UK ONLY\][^]]*', '', text, flags=re.IGNORECASE)
            text = re.sub(r'\[FOR JP ONLY\][^]]*', '', text, flags=re.IGNORECASE)
            text = re.sub(r'\[SHOW HYPERLINK[^\]]*\]', '', text, flags=re.IGNORECASE)
            text = re.sub(r'\[RANDOMIZE[^\]]*\]', '', text, flags=re.IGNORECASE)
            text = re.sub(r'\[CAPTURE[^\]]*\]', '', text, flags=re.IGNORECASE)
            
            # Clean up whitespace
            text = re.sub(r'\s+', ' ', text).strip()
            question['text'] = text
            
            # Ensure standard yes/no options if not present
            if 'options' not in question or not question['options']:
                question['options'] = ['Yes', 'No']
            elif len(question['options']) == 2:
                # Clean up existing options
                cleaned_options = []
                for option in question['options']:
                    clean_option = re.sub(r'\[.*?\]', '', str(option)).strip()
                    if clean_option:
                        cleaned_options.append(clean_option)
                question['options'] = cleaned_options
            
            logger.info(f"Successfully post-processed yes_no question {question.get('id')}")
            return question
            
        except Exception as e:
            logger.error(f"Error post-processing yes_no {question.get('id')}: {e}", exc_info=True)
            return original_question

    def _post_process_text_block(self, text_block: Dict[str, Any]) -> Dict[str, Any]:
        """Clean text blocks - ensure proper structure and remove programming instructions."""
        import copy
        import re
        
        original_text_block = copy.deepcopy(text_block)
        
        try:
            # Handle different text block structures
            text_content = text_block.get('text') or text_block.get('content', '')
            
            if not text_content:
                return original_text_block
            
            # Remove programming instructions
            text_content = re.sub(r'\[FOR US ONLY\][^]]*', '', text_content, flags=re.IGNORECASE)
            text_content = re.sub(r'\[FOR UK ONLY\][^]]*', '', text_content, flags=re.IGNORECASE)
            text_content = re.sub(r'\[FOR JP ONLY\][^]]*', '', text_content, flags=re.IGNORECASE)
            text_content = re.sub(r'\[SHOW HYPERLINK[^\]]*\]', '', text_content, flags=re.IGNORECASE)
            text_content = re.sub(r'\[RANDOMIZE[^\]]*\]', '', text_content, flags=re.IGNORECASE)
            text_content = re.sub(r'\[CAPTURE[^\]]*\]', '', text_content, flags=re.IGNORECASE)
            
            # Clean up whitespace
            text_content = re.sub(r'\s+', ' ', text_content).strip()
            
            # Update the text content
            if 'text' in text_block:
                text_block['text'] = text_content
            if 'content' in text_block:
                text_block['content'] = text_content
            
            # Ensure we have a text property for the UI
            if 'text' not in text_block and 'content' in text_block:
                text_block['text'] = text_block['content']
            
            logger.info(f"Successfully post-processed text block")
            return text_block
            
        except Exception as e:
            logger.error(f"Error post-processing text block: {e}", exc_info=True)
            return original_text_block
    
    async def convert_to_json(self, document_text: str, comments: List[Dict[str, Any]] = None, session_id: str = None) -> Dict[str, Any]:
        """Convert document text to JSON using LLM with detailed progress updates."""
        logger.info(f"🤖 [Document Parser] Starting LLM conversion with model: {self.model}")
        try:
            # Send initial progress
            if session_id:
                await self._send_progress(session_id, "llm_processing", 30,
                    "Preparing AI analysis...",
                    "Building conversion prompt and initializing AI model",
                    estimated_time=45)
            
            logger.info(f"📝 [Document Parser] Creating conversion prompt")
            prompt = self.create_conversion_prompt(document_text, comments)
            logger.info(f"✅ [Document Parser] Conversion prompt created, length: {len(prompt)} chars")
            
            if session_id:
                await self._send_progress(session_id, "llm_processing", 35,
                    "Sending to AI model...",
                    f"Using {self.model} to analyze document structure",
                    estimated_time=40)
            
            # Create audit context for this LLM interaction
            interaction_id = f"document_parsing_{uuid.uuid4().hex[:8]}"
            audit_service = LLMAuditService(self.db_session) if self.db_session else None
            
            # Try to use audit service, but don't let audit failures break the core functionality
            try:
                if audit_service:
                    async with LLMAuditContext(
                        audit_service=audit_service,
                        interaction_id=interaction_id,
                        model_name=self.model,
                        model_provider="replicate",
                        purpose="document_parsing",
                        input_prompt=prompt,
                        sub_purpose="survey_conversion",
                        context_type="document",
                        hyperparameters={
                            "temperature": 0.1,
                            "max_tokens": 4000
                        },
                        metadata={
                            "document_length": len(document_text),
                            "prompt_length": len(prompt)
                        },
                        tags=["document_parsing", "survey_conversion"]
                    ) as audit_context:
                        logger.info(f"🚀 [Document Parser] Calling Replicate API with auditing")
                        logger.info(f"🎯 [Document Parser] Using model: {self.model}")
                        logger.info(f"🎯 [Document Parser] Model source: DocumentParser._parse_document_with_llm (audited)")
                        
                        if session_id:
                            await self._send_progress(session_id, "llm_processing", 40,
                                "AI processing document...",
                                "Analyzing survey structure and extracting questions",
                                estimated_time=35)
                        
                        start_time = time.time()
                        output = await self.replicate_client.async_run(
                            self.model,
                            input={
                                "prompt": prompt,
                                "temperature": 0.1,
                                "max_tokens": 2000,  # Reduced from 4000
                                "system_prompt": "Parse document to JSON. Return ONLY valid JSON with this structure:\n\n{\n  \"raw_output\": {\n    \"document_text\": \"[original text]\",\n    \"extraction_timestamp\": \"2024-01-01T00:00:00Z\",\n    \"source_file\": null,\n    \"error\": null\n  },\n  \"final_output\": {\n    \"title\": \"[survey title]\",\n    \"description\": \"[description or null]\",\n    \"metadata\": {\n      \"quality_score\": 0.9,\n      \"estimated_time\": 10,\n      \"methodology_tags\": [\"tag1\", \"tag2\"],\n      \"target_responses\": 100,\n      \"source_file\": null\n    },\n    \"sections\": [\n      {\n        \"id\": 1,\n        \"title\": \"[section title]\",\n        \"description\": \"[section description or null]\",\n        \"introText\": {\n          \"text\": \"[intro text or null]\"\n        },\n        \"textBlocks\": [\n          {\n            \"text\": \"[additional text block or null]\"\n          }\n        ],\n        \"questions\": [\n          {\n            \"id\": \"q1\",\n            \"text\": \"[question text]\",\n            \"type\": \"multiple_choice|single_choice|text|scale|rating|yes_no|dropdown|matrix|ranking|numeric|date|boolean|file_upload|van_westendorp|gabor_granger|conjoint|maxdiff|instruction|display_only|unknown\",\n            \"options\": [\"option1\", \"option2\"],\n            \"required\": true,\n            \"validation\": \"single_select\",\n            \"methodology\": \"van_westendorp|gabor_granger|conjoint|maxdiff|null\",\n            \"routing\": null,\n            \"order\": 1\n          }\n        ]\n      }\n    ],\n    \"parsing_issues\": []\n  }\n}\n\nCRITICAL RULES:\n- MUST use sections format with 7 sections (Sample Plan, Screener, Brand/Product Awareness, Concept Exposure, Methodology, Additional Questions, Programmer Instructions)\n- DO NOT use flat 'questions' array - use 'sections' array instead\n- Each section has id (1-7), title, description, introText, textBlocks, and questions array\n- Use sequential question IDs: q1, q2, q3...\n- For choice questions: put options in \"options\" array\n- For scales: use \"type\":\"scale\" with scale labels in \"options\"\n- For matrices: use \"type\":\"matrix\"\n- Set \"required\": true for most questions\n- Use \"instruction\" type for non-interactive text blocks\n- Return ONLY JSON, no explanations\n\nMethodology Detection:\n- Van Westendorp: \"too cheap\", \"too expensive\" price questions\n- Gabor-Granger: Sequential price acceptance\n- Conjoint: Choice scenarios with attributes\n- MaxDiff: \"Most/Least important\" selections\n- NPS: \"How likely to recommend\" questions"
                            }
                        )
                        
                        # Process the output and set audit context
                        response_time_ms = int((time.time() - start_time) * 1000)
                        
                        # Capture raw response immediately (unprocessed)
                        if hasattr(output, '__iter__') and not isinstance(output, str):
                            raw_response = "".join(str(chunk) for chunk in output)
                        else:
                            raw_response = str(output)
                        
                        # Process the output for further use
                        processed_output = raw_response.strip()
                        
                        # Set raw response (unprocessed) and processed output
                        audit_context.set_raw_response(raw_response)
                        audit_context.set_output(
                            output_content=processed_output
                        )
                else:
                    # Fallback without auditing
                    logger.info(f"🚀 [Document Parser] Calling Replicate API without auditing")
                    logger.info(f"🎯 [Document Parser] Using model: {self.model}")
                    logger.info(f"🎯 [Document Parser] Model source: DocumentParser._parse_document_with_llm (non-audited)")
                    output = await self.replicate_client.async_run(
                        self.model,
                        input={
                        "prompt": prompt,
                        "temperature": 0.1,
                        "max_tokens": 2000,  # Reduced from 4000
                        "system_prompt": "Parse document to JSON. Return ONLY valid JSON with this structure:\n\n{\n  \"raw_output\": {\n    \"document_text\": \"[original text]\",\n    \"extraction_timestamp\": \"2024-01-01T00:00:00Z\",\n    \"source_file\": null,\n    \"error\": null\n  },\n  \"final_output\": {\n    \"title\": \"[survey title]\",\n    \"description\": \"[description or null]\",\n    \"metadata\": {\n      \"quality_score\": 0.9,\n      \"estimated_time\": 10,\n      \"methodology_tags\": [\"tag1\", \"tag2\"],\n      \"target_responses\": 100,\n      \"source_file\": null\n    },\n    \"sections\": [\n      {\n        \"id\": 1,\n        \"title\": \"[section title]\",\n        \"description\": \"[section description or null]\",\n        \"introText\": {\n          \"text\": \"[intro text or null]\"\n        },\n        \"textBlocks\": [\n          {\n            \"text\": \"[additional text block or null]\"\n          }\n        ],\n        \"questions\": [\n          {\n            \"id\": \"q1\",\n            \"text\": \"[question text]\",\n            \"type\": \"multiple_choice|single_choice|text|scale|rating|yes_no|dropdown|matrix|ranking|numeric|date|boolean|file_upload|van_westendorp|gabor_granger|conjoint|maxdiff|instruction|display_only|unknown\",\n            \"options\": [\"option1\", \"option2\"],\n            \"required\": true,\n            \"validation\": \"single_select\",\n            \"methodology\": \"van_westendorp|gabor_granger|conjoint|maxdiff|null\",\n            \"routing\": null,\n            \"order\": 1\n          }\n        ]\n      }\n    ],\n    \"parsing_issues\": []\n  }\n}\n\nCRITICAL RULES:\n- MUST use sections format with 7 sections (Sample Plan, Screener, Brand/Product Awareness, Concept Exposure, Methodology, Additional Questions, Programmer Instructions)\n- DO NOT use flat 'questions' array - use 'sections' array instead\n- Each section has id (1-7), title, description, introText, textBlocks, and questions array\n- Use sequential question IDs: q1, q2, q3...\n- For choice questions: put options in \"options\" array\n- For scales: use \"type\":\"scale\" with scale labels in \"options\"\n- For matrices: use \"type\":\"matrix\"\n- Set \"required\": true for most questions\n- Use \"instruction\" type for non-interactive text blocks\n- Return ONLY JSON, no explanations\n\nMethodology Detection:\n- Van Westendorp: \"too cheap\", \"too expensive\" price questions\n- Gabor-Granger: Sequential price acceptance\n- Conjoint: Choice scenarios with attributes\n- MaxDiff: \"Most/Least important\" selections\n- NPS: \"How likely to recommend\" questions"
                    }
                )
            except Exception as audit_error:
                # If audit fails, log the error but continue with core functionality
                logger.warning(f"⚠️ [Document Parser] Audit system failed, continuing without audit: {str(audit_error)}")
                logger.info(f"🚀 [Document Parser] Calling Replicate API without auditing (audit failed)")
                output = await self.replicate_client.async_run(
                    self.model,
                    input={
                        "prompt": prompt,
                        "temperature": 0.1,
                        "max_tokens": 2000,  # Reduced from 4000
                        "system_prompt": "Parse document to JSON. Return ONLY valid JSON with this structure:\n\n{\n  \"raw_output\": {\n    \"document_text\": \"[original text]\",\n    \"extraction_timestamp\": \"2024-01-01T00:00:00Z\",\n    \"source_file\": null,\n    \"error\": null\n  },\n  \"final_output\": {\n    \"title\": \"[survey title]\",\n    \"description\": \"[description or null]\",\n    \"metadata\": {\n      \"quality_score\": 0.9,\n      \"estimated_time\": 10,\n      \"methodology_tags\": [\"tag1\", \"tag2\"],\n      \"target_responses\": 100,\n      \"source_file\": null\n    },\n    \"sections\": [\n      {\n        \"id\": 1,\n        \"title\": \"[section title]\",\n        \"description\": \"[section description or null]\",\n        \"introText\": {\n          \"text\": \"[intro text or null]\"\n        },\n        \"textBlocks\": [\n          {\n            \"text\": \"[additional text block or null]\"\n          }\n        ],\n        \"questions\": [\n          {\n            \"id\": \"q1\",\n            \"text\": \"[question text]\",\n            \"type\": \"multiple_choice|single_choice|text|scale|rating|yes_no|dropdown|matrix|ranking|numeric|date|boolean|file_upload|van_westendorp|gabor_granger|conjoint|maxdiff|instruction|display_only|unknown\",\n            \"options\": [\"option1\", \"option2\"],\n            \"required\": true,\n            \"validation\": \"single_select\",\n            \"methodology\": \"van_westendorp|gabor_granger|conjoint|maxdiff|null\",\n            \"routing\": null,\n            \"order\": 1\n          }\n        ]\n      }\n    ],\n    \"parsing_issues\": []\n  }\n}\n\nCRITICAL RULES:\n- MUST use sections format with 7 sections (Sample Plan, Screener, Brand/Product Awareness, Concept Exposure, Methodology, Additional Questions, Programmer Instructions)\n- DO NOT use flat 'questions' array - use 'sections' array instead\n- Each section has id (1-7), title, description, introText, textBlocks, and questions array\n- Use sequential question IDs: q1, q2, q3...\n- For choice questions: put options in \"options\" array\n- For scales: use \"type\":\"scale\" with scale labels in \"options\"\n- For matrices: use \"type\":\"matrix\"\n- Set \"required\": true for most questions\n- Use \"instruction\" type for non-interactive text blocks\n- Return ONLY JSON, no explanations\n\nMethodology Detection:\n- Van Westendorp: \"too cheap\", \"too expensive\" price questions\n- Gabor-Granger: Sequential price acceptance\n- Conjoint: Choice scenarios with attributes\n- MaxDiff: \"Most/Least important\" selections\n- NPS: \"How likely to recommend\" questions"
                    }
                )
            
            # Replicate returns a generator, join the output
            logger.info(f"📥 [Document Parser] Processing LLM response")
            logger.debug(f"📥 [Document Parser] Output type: {type(output)}")
            logger.debug(f"📥 [Document Parser] Output value: {output}")
            
            # Handle different output types from Replicate
            if hasattr(output, '__iter__') and not isinstance(output, str):
                json_content = "".join(str(chunk) for chunk in output).strip()
            else:
                json_content = str(output).strip()
            
            # CRITICAL FIX: Handle character array output from LLM
            # Sometimes the LLM returns a character array instead of a JSON string
            logger.info(f"🔍 [Document Parser] Checking for character array format...")
            logger.info(f"🔍 [Document Parser] Content starts with: {json_content[:50]}")
            logger.info(f"🔍 [Document Parser] Content ends with: {json_content[-50:]}")
            
            if json_content.startswith('[') and json_content.endswith(']'):
                logger.info(f"🔧 [Document Parser] Detected potential character array format")
                try:
                    # Try to parse as a character array and join it
                    import ast
                    char_array = ast.literal_eval(json_content)
                    if isinstance(char_array, list) and all(isinstance(c, str) for c in char_array):
                        original_length = len(json_content)
                        json_content = ''.join(char_array).strip()
                        logger.info(f"🔧 [Document Parser] Successfully converted character array to string")
                        logger.info(f"🔧 [Document Parser] Original length: {original_length}, New length: {len(json_content)}")
                        logger.info(f"🔧 [Document Parser] First 200 chars: {json_content[:200]}")
                    else:
                        logger.warning(f"⚠️ [Document Parser] Character array contains non-string elements")
                        logger.warning(f"⚠️ [Document Parser] Array type: {type(char_array)}, First few elements: {char_array[:5] if char_array else 'empty'}")
                except Exception as e:
                    logger.warning(f"⚠️ [Document Parser] Failed to parse character array: {e}")
                    logger.warning(f"⚠️ [Document Parser] Raw content: {json_content[:100]}")
            else:
                logger.info(f"🔍 [Document Parser] Content does not appear to be a character array")
            
            logger.info(f"📝 [Document Parser] Final JSON content length: {len(json_content)}")
            logger.debug(f"📝 [Document Parser] JSON content preview: {json_content[:200]}...")
                
            logger.info(f"✅ [Document Parser] LLM response received, length: {len(json_content)} chars")
            logger.info(f"📄 [Document Parser] LLM response preview: {json_content[:500]}...")
            logger.info(f"📄 [Document Parser] LLM response ending: ...{json_content[-200:]}")
            
            # Try to parse the JSON
            logger.info(f"🔍 [Document Parser] Parsing JSON response")
            logger.info(f"🔍 [Document Parser] JSON content length: {len(json_content)}")
            logger.info(f"🔍 [Document Parser] JSON content starts with: {json_content[:50]}")
            logger.info(f"🔍 [Document Parser] JSON content ends with: {json_content[-50:]}")
            
            # Use centralized JSON parsing utility
            survey_data = parse_llm_json_response(json_content, service_name="DocumentParser")
            
            # If parsing failed, try to extract JSON from text (for character array responses)
            if survey_data is None:
                logger.info(f"🔧 [Document Parser] Standard JSON parsing failed, trying JSON extraction...")
                import re
                
                # Try to find JSON between ``` markers
                json_match = re.search(r'```(?:json)?\s*(\{.*?\})\s*```', json_content, re.DOTALL)
                if json_match:
                    json_text = json_match.group(1)
                    logger.info(f"🔧 [Document Parser] Found JSON between ``` markers")
                    survey_data = parse_llm_json_response(json_text, service_name="DocumentParser")
                else:
                    # Try to find JSON object starting with {
                    json_match = re.search(r'(\{.*\})', json_content, re.DOTALL)
                    if json_match:
                        json_text = json_match.group(1)
                        logger.info(f"🔧 [Document Parser] Found JSON object in text")
                        survey_data = parse_llm_json_response(json_text, service_name="DocumentParser")
                    else:
                        logger.warning(f"⚠️ [Document Parser] No JSON object found in converted text")
            
            if survey_data is None:
                logger.error(f"❌ [Document Parser] All JSON extraction methods failed")
                logger.error(f"❌ [Document Parser] Raw response (first 1000 chars): {json_content[:1000]}")
                raise DocumentParsingError(f"No valid JSON found in response")
            
            logger.info(f"✅ [Document Parser] JSON parsing successful")
            logger.info(f"📊 [Document Parser] Parsed data keys: {list(survey_data.keys()) if isinstance(survey_data, dict) else 'Not a dict'}")
            logger.info(f"📊 [Document Parser] Full parsed data: {survey_data}")
            
            # Validate the expected structure
            if not isinstance(survey_data, dict):
                raise ValueError("Response is not a JSON object")
            
            if "raw_output" not in survey_data or "final_output" not in survey_data:
                logger.warning(f"⚠️ [Document Parser] Missing expected structure (raw_output/final_output)")
                
                # Add comprehensive logging of response structure
                logger.info(f"📊 [Document Parser] Response structure keys: {list(survey_data.keys())}")
                if isinstance(survey_data, dict):
                    logger.info(f"📊 [Document Parser] Has 'sections': {'sections' in survey_data}")
                    logger.info(f"📊 [Document Parser] Has 'questions': {'questions' in survey_data}")
                    logger.info(f"📊 [Document Parser] Has 'title': {'title' in survey_data}")
                
                # Try to use the response as-is if it has the old structure
                if "title" in survey_data and ("questions" in survey_data or "sections" in survey_data):
                    # Validate structure quality before wrapping
                    if not self._validate_llm_response_quality(survey_data):
                        raise ValueError("LLM response failed quality validation")
                    
                    if "sections" in survey_data:
                        logger.info(f"🔧 [Document Parser] Using sections format with {len(survey_data['sections'])} sections")
                    elif "questions" in survey_data:
                        logger.info(f"🔧 [Document Parser] Using legacy questions format with {len(survey_data['questions'])} questions")
                    
                    survey_data = {
                        "raw_output": {"document_text": document_text, "extraction_timestamp": "2024-01-01T00:00:00Z"},
                        "final_output": survey_data
                    }
                else:
                    raise ValueError("Response missing required raw_output and final_output structure")
            
            logger.info(f"✅ [Document Parser] JSON structure validation successful")
            
            logger.info(f"🎉 [Document Parser] LLM conversion completed successfully")
            return survey_data
            
        except Exception as e:
            logger.error(f"❌ [Document Parser] Failed to convert document to JSON: {str(e)}", exc_info=True)
            
            # Create a fallback response with error details
            fallback_response = {
                "raw_output": {
                    "document_text": document_text,
                    "extraction_timestamp": "2024-01-01T00:00:00Z",
                    "source_file": None,
                    "error": f"LLM conversion failed: {str(e)}"
                },
                "final_output": {
                    "title": "Document Parse Error",
                    "description": f"Failed to parse document: {str(e)}",
                    "metadata": {
                        "quality_score": 0.0,
                        "estimated_time": 0,
                        "methodology_tags": [],
                        "target_responses": None,
                        "source_file": None
                    },
                    "questions": [],
                    "parsing_issues": [f"LLM conversion failed: {str(e)}"]
                }
            }
            
            logger.warning(f"⚠️ [Document Parser] Returning fallback response due to error")
            return fallback_response
    
    def _validate_llm_response_quality(self, survey_data: Dict[str, Any]) -> bool:
        """
        Validate LLM response has minimum required quality.
        Uses lenient validation - logs warnings but allows empty sections.
        """
        if not isinstance(survey_data, dict):
            logger.error("❌ [Document Parser] Response is not a dictionary")
            return False
        
        if "title" not in survey_data or not survey_data.get("title"):
            logger.warning("⚠️ [Document Parser] Missing or empty title (allowed but not ideal)")
        
        # Check sections format
        if "sections" in survey_data:
            sections = survey_data["sections"]
            if not isinstance(sections, list):
                logger.error("❌ [Document Parser] Sections is not a list")
                return False
            
            if len(sections) == 0:
                logger.warning("⚠️ [Document Parser] Sections array is empty (allowed)")
            else:
                # Count sections with questions
                sections_with_questions = sum(
                    1 for s in sections
                    if isinstance(s, dict) and isinstance(s.get("questions"), list) and len(s["questions"]) > 0
                )
                logger.info(f"📊 [Document Parser] {sections_with_questions}/{len(sections)} sections have questions")
                
                if sections_with_questions == 0:
                    logger.warning("⚠️ [Document Parser] No sections contain questions (allowed but suspicious)")
        
        # Check legacy format
        elif "questions" in survey_data:
            questions = survey_data["questions"]
            if not isinstance(questions, list):
                logger.error("❌ [Document Parser] Questions is not a list")
                return False
            
            if len(questions) == 0:
                logger.warning("⚠️ [Document Parser] Questions array is empty (allowed)")
            else:
                logger.info(f"📊 [Document Parser] Found {len(questions)} questions in legacy format")
        else:
            logger.error("❌ [Document Parser] No sections or questions found in response")
            return False
        
        return True
    
    def validate_survey_json(self, survey_data: Dict[str, Any]) -> Dict[str, Any]:
        """Validate the generated JSON against our survey schema."""
        try:
            # Extract the final_output for validation
            if "final_output" in survey_data:
                final_output = survey_data["final_output"]
                logger.info(f"🔍 [Document Parser] Validating final_output structure")
                
                # Handle sections format (NEW)
                if "sections" in final_output and isinstance(final_output["sections"], list):
                    logger.info(f"🔍 [Document Parser] Processing sections format with {len(final_output['sections'])} sections")
                    
                    # Process questions in each section
                    for section_idx, section in enumerate(final_output["sections"]):
                        if isinstance(section, dict) and "questions" in section and isinstance(section["questions"], list):
                            logger.info(f"🔍 [Document Parser] Processing {len(section['questions'])} questions in section {section.get('id', section_idx+1)}")
                            
                            for i, question in enumerate(section["questions"]):
                                if isinstance(question, dict) and "order" not in question:
                                    question["order"] = i + 1
                                
                                # Fix scale_labels format - convert array to dictionary if needed
                                if isinstance(question, dict) and "scale_labels" in question:
                                    scale_labels = question["scale_labels"]
                                    if isinstance(scale_labels, list) and len(scale_labels) >= 2:
                                        # Convert array to dictionary format
                                        question["scale_labels"] = {
                                            "min": scale_labels[0],
                                            "max": scale_labels[-1]
                                        }
                                        logger.debug(f"🔧 [Document Parser] Fixed scale_labels for section {section.get('id', section_idx+1)} question {i+1}: {scale_labels} -> {question['scale_labels']}")
                                    elif not isinstance(scale_labels, dict):
                                        # Remove invalid scale_labels
                                        question.pop("scale_labels", None)
                                        logger.debug(f"🔧 [Document Parser] Removed invalid scale_labels for section {section.get('id', section_idx+1)} question {i+1}")
                
                # Handle legacy flat format (for backward compatibility)
                elif "questions" in final_output and isinstance(final_output["questions"], list):
                    logger.info(f"🔍 [Document Parser] Processing legacy flat format with {len(final_output['questions'])} questions")
                    
                    for i, question in enumerate(final_output["questions"]):
                        if isinstance(question, dict) and "order" not in question:
                            question["order"] = i + 1
                        
                        # Fix scale_labels format - convert array to dictionary if needed
                        if isinstance(question, dict) and "scale_labels" in question:
                            scale_labels = question["scale_labels"]
                            if isinstance(scale_labels, list) and len(scale_labels) >= 2:
                                # Convert array to dictionary format
                                question["scale_labels"] = {
                                    "min": scale_labels[0],
                                    "max": scale_labels[-1]
                                }
                                logger.debug(f"🔧 [Document Parser] Fixed scale_labels for question {i+1}: {scale_labels} -> {question['scale_labels']}")
                            elif not isinstance(scale_labels, dict):
                                # Remove invalid scale_labels
                                question.pop("scale_labels", None)
                                logger.debug(f"🔧 [Document Parser] Removed invalid scale_labels for question {i+1}")
                
                # Extract annotation fields before validation (SurveyCreate doesn't support them)
                annotation_fields = {}
                
                # Handle sections format for annotations
                if "sections" in final_output and isinstance(final_output["sections"], list):
                    for section_idx, section in enumerate(final_output["sections"]):
                        if isinstance(section, dict) and "questions" in section and isinstance(section["questions"], list):
                            for i, question in enumerate(section["questions"]):
                                if isinstance(question, dict) and "annotation" in question:
                                    annotation_fields[question.get("id", f"s{section.get('id', section_idx+1)}_q{i+1}")] = question.pop("annotation")
                
                # Handle legacy format for annotations
                elif "questions" in final_output and isinstance(final_output["questions"], list):
                    for i, question in enumerate(final_output["questions"]):
                        if isinstance(question, dict) and "annotation" in question:
                            annotation_fields[question.get("id", f"q{i+1}")] = question.pop("annotation")
                
                # Create appropriate SurveyCreate object based on structure
                if "sections" in final_output:
                    logger.info(f"✅ [Document Parser] Using SurveyCreateWithSections for validation")
                    # Use sections format validation
                    survey = SurveyCreateWithSections(**final_output)
                else:
                    logger.info(f"✅ [Document Parser] Using SurveyCreate (legacy) for validation")
                    # Use legacy questions format validation
                    survey = SurveyCreate(**final_output)
                validated_data = survey.model_dump()
                
                # Restore annotation fields after validation
                if annotation_fields:
                    # Handle sections format for restoring annotations
                    if "sections" in validated_data and isinstance(validated_data["sections"], list):
                        for section in validated_data["sections"]:
                            if isinstance(section, dict) and "questions" in section and isinstance(section["questions"], list):
                                for question in section["questions"]:
                                    question_id = question.get("id")
                                    if question_id in annotation_fields:
                                        question["annotation"] = annotation_fields[question_id]
                    
                    # Handle legacy format for restoring annotations
                    elif "questions" in validated_data and isinstance(validated_data["questions"], list):
                        for question in validated_data["questions"]:
                            question_id = question.get("id")
                            if question_id in annotation_fields:
                                question["annotation"] = annotation_fields[question_id]
                
                # Preserve metadata fields that are not part of SurveyCreate model
                metadata_fields = {}
                for field in ['product_category', 'research_goal', 'methodologies', 'target_segment', 'estimated_time', 'confidence_score']:
                    if field in final_output:
                        metadata_fields[field] = final_output[field]
                
                # Add metadata fields to validated_data
                validated_data.update(metadata_fields)
                
                # Return the full structure with validated final_output
                return {
                    "raw_output": survey_data.get("raw_output", {}),
                    "final_output": validated_data
                }
            else:
                # Fallback for legacy format
                logger.warning(f"⚠️ [Document Parser] Using legacy format validation")
                
                # Fix scale_labels format for legacy format
                if "questions" in survey_data and isinstance(survey_data["questions"], list):
                    for i, question in enumerate(survey_data["questions"]):
                        if isinstance(question, dict) and "scale_labels" in question:
                            scale_labels = question["scale_labels"]
                            if isinstance(scale_labels, list) and len(scale_labels) >= 2:
                                # Convert array to dictionary format
                                question["scale_labels"] = {
                                    "min": scale_labels[0],
                                    "max": scale_labels[-1]
                                }
                                logger.debug(f"🔧 [Document Parser] Fixed scale_labels for legacy question {i+1}: {scale_labels} -> {question['scale_labels']}")
                            elif not isinstance(scale_labels, dict):
                                # Remove invalid scale_labels
                                question.pop("scale_labels", None)
                                logger.debug(f"🔧 [Document Parser] Removed invalid scale_labels for legacy question {i+1}")
                
                # Use appropriate validation based on structure
                if "sections" in survey_data and "questions" not in survey_data:
                    survey = SurveyCreateWithSections(**survey_data)
                else:
                    survey = SurveyCreate(**survey_data)
                validated_data = survey.model_dump()
                
                # Preserve metadata fields that are not part of SurveyCreate model
                metadata_fields = {}
                for field in ['product_category', 'research_goal', 'methodologies', 'target_segment', 'estimated_time', 'confidence_score']:
                    if field in survey_data:
                        metadata_fields[field] = survey_data[field]
                
                # Add metadata fields to validated_data
                validated_data.update(metadata_fields)
                
                return validated_data
                
        except ValidationError as e:
            logger.error(f"Survey validation failed: {str(e)}")
            logger.error(f"Raw data structure: {list(survey_data.keys()) if isinstance(survey_data, dict) else 'Not a dict'}")
            if "final_output" in survey_data:
                logger.error(f"Final output structure: {list(survey_data['final_output'].keys()) if isinstance(survey_data.get('final_output'), dict) else 'Not a dict'}")
            
            # Try to fix invalid question types first
            try:
                logger.warning(f"⚠️ [Document Parser] Attempting to fix invalid question types")
                fixed_data = self._fix_invalid_question_types(survey_data)
                
                if "final_output" in fixed_data:
                    survey = SurveyCreate(**fixed_data["final_output"])
                    validated_data = survey.model_dump()
                    logger.info(f"✅ [Document Parser] Survey validation successful after fixing question types")
                    return {
                        "raw_output": fixed_data.get("raw_output", {}),
                        "final_output": validated_data
                    }
                else:
                    survey = SurveyCreate(**fixed_data)
                    logger.info(f"✅ [Document Parser] Survey validation successful after fixing question types")
                    return survey.model_dump()
            except Exception as fix_error:
                logger.warning(f"⚠️ [Document Parser] Question type fixing failed: {str(fix_error)}")
            
            # Try to fix routing fields
            try:
                logger.warning(f"⚠️ [Document Parser] Attempting to fix routing fields")
                fixed_data = self._fix_routing_fields(survey_data)
                
                if "final_output" in fixed_data:
                    survey = SurveyCreate(**fixed_data["final_output"])
                    validated_data = survey.model_dump()
                    logger.info(f"✅ [Document Parser] Survey validation successful after fixing routing fields")
                    return {
                        "raw_output": fixed_data.get("raw_output", {}),
                        "final_output": validated_data
                    }
                else:
                    survey = SurveyCreate(**fixed_data)
                    logger.info(f"✅ [Document Parser] Survey validation successful after fixing routing fields")
                    return survey.model_dump()
            except Exception as fix_error:
                logger.warning(f"⚠️ [Document Parser] Routing field fixing failed: {str(fix_error)}")
            
            # Try to create a minimal valid response
            try:
                logger.warning(f"⚠️ [Document Parser] Attempting to create minimal valid response")
                minimal_final_output = {
                    "title": survey_data.get("final_output", {}).get("title", "Untitled Survey"),
                    "description": survey_data.get("final_output", {}).get("description", "Survey description not available"),
                    "metadata": {
                        "quality_score": 0.3,
                        "estimated_time": 5,
                        "methodology_tags": [],
                        "target_responses": None,
                        "source_file": None
                    },
                    "questions": survey_data.get("final_output", {}).get("questions", []),
                    "parsing_issues": [f"Validation failed: {str(e)}"]
                }
                
                # Validate the minimal response
                survey = SurveyCreate(**minimal_final_output)
                validated_data = survey.model_dump()
                
                return {
                    "raw_output": survey_data.get("raw_output", {}),
                    "final_output": validated_data
                }
                
            except Exception as e2:
                logger.error(f"❌ [Document Parser] Minimal response creation also failed: {str(e2)}")
                raise DocumentParsingError(f"Generated JSON validation failed: {str(e)}")
    
    def _infer_question_type_from_content(self, question: Dict[str, Any]) -> str:
        """Intelligently infer question type from content when type is unknown."""
        text = question.get("text", "").lower()
        options = question.get("options", [])
        scale_labels = question.get("scale_labels", {})
        
        # If it has options, it's likely a choice question
        if options and len(options) > 0:
            if len(options) == 2 and any(opt.lower() in ['yes', 'no', 'true', 'false'] for opt in options):
                return 'yes_no'
            elif len(options) > 10:  # Many options suggest dropdown
                return 'dropdown'
            else:
                return 'single_choice'
        
        # If it has scale labels, it's a scale question
        if scale_labels:
            return 'scale'
        
        # Check for instructional/informational content
        instructional_keywords = [
            'click', 'continue', 'review', 'introduction', 'thank you', 'welcome',
            'please note', 'disclaimer', 'instruction', 'guidance', 'proceed',
            'next', 'begin', 'start', 'complete', 'finish', 'end'
        ]
        
        if any(keyword in text for keyword in instructional_keywords):
            return 'instruction'
        
        # Check for matrix/grid patterns
        if 'matrix' in text or 'grid' in text or 'table' in text:
            return 'matrix'
        
        # Default to text for open-ended questions
        return 'text'

    def _fix_invalid_question_types(self, survey_data: Dict[str, Any]) -> Dict[str, Any]:
        """Fix invalid question types by mapping them to valid ones."""
        logger.info(f"🔧 [Document Parser] Attempting to fix invalid question types")
        
        # Mapping of invalid types to valid ones
        type_mapping = {
            'numeric': 'text',
            'number': 'text', 
            'integer': 'text',
            'float': 'text',
            'decimal': 'text',
            'currency': 'text',
            'email': 'text',
            'phone': 'text',
            'date': 'text',
            'time': 'text',
            'datetime': 'text',
            'url': 'text',
            'website': 'text',
            'address': 'text',
            'name': 'text',
            'age': 'text',
            'gender': 'single_choice',
            'yes_no': 'yes_no',
            'boolean': 'yes_no',
            'true_false': 'yes_no',
            'agree_disagree': 'scale',
            'likert': 'scale',
            'rating': 'scale',
            'stars': 'scale',
            'slider': 'scale',
            'range': 'scale',
            'dropdown': 'dropdown',
            'select': 'single_choice',
            'radio': 'single_choice',
            'checkbox': 'multiple_choice',
            'multi_select': 'multiple_choice',
            'checkboxes': 'multiple_choice',
            'ranking': 'ranking',
            'rank': 'ranking',
            'order': 'ranking',
            'priority': 'ranking',
            'matrix': 'matrix',
            'grid': 'matrix',
            'table': 'matrix',
            'open_text': 'text',
            'open_ended': 'text',
            'comment': 'text',
            'feedback': 'text',
            'suggestion': 'text',
            'other': 'text',
            'custom': 'text',
            'unknown': 'unknown'  # Will be handled by intelligent inference
        }
        
        fixed_data = survey_data.copy()
        
        # Handle both new format (with final_output) and legacy format
        if "final_output" in fixed_data:
            final_output = fixed_data["final_output"].copy()
            
            # Handle sections format (NEW)
            if 'sections' in final_output and isinstance(final_output['sections'], list):
                logger.info(f"🔧 [Document Parser] Fixing question types in sections format with {len(final_output['sections'])} sections")
                
                for section_idx, section in enumerate(final_output['sections']):
                    if isinstance(section, dict) and 'questions' in section and isinstance(section['questions'], list):
                        fixed_questions = []
                        for question in section['questions']:
                            fixed_question = question.copy()
                            if 'type' in fixed_question:
                                original_type = fixed_question['type']
                                if original_type not in [t.value for t in QuestionType]:
                                    mapped_type = type_mapping.get(original_type.lower(), 'unknown')
                                    logger.info(f"🔧 [Document Parser] Mapped question type '{original_type}' to '{mapped_type}' in section {section.get('id', section_idx+1)}")
                                    
                                    # If still unknown, use intelligent inference
                                    if mapped_type == 'unknown':
                                        inferred_type = self._infer_question_type_from_content(fixed_question)
                                        logger.info(f"🔧 [Document Parser] Intelligently inferred type '{inferred_type}' for question: {fixed_question.get('text', '')[:50]}...")
                                        mapped_type = inferred_type
                                    
                                    fixed_question['type'] = mapped_type
                            
                            # Fix scale_labels format
                            if 'scale_labels' in fixed_question:
                                scale_labels = fixed_question['scale_labels']
                                if isinstance(scale_labels, list) and len(scale_labels) >= 2:
                                    fixed_question['scale_labels'] = {
                                        "min": scale_labels[0],
                                        "max": scale_labels[-1]
                                    }
                                elif not isinstance(scale_labels, dict):
                                    fixed_question.pop('scale_labels', None)
                            
                            fixed_questions.append(fixed_question)
                        section['questions'] = fixed_questions
                        logger.info(f"🔧 [Document Parser] Fixed {len(fixed_questions)} questions in section {section.get('id', section_idx+1)}")
            
            # Handle legacy flat format (for backward compatibility)
            elif 'questions' in final_output and isinstance(final_output['questions'], list):
                logger.info(f"🔧 [Document Parser] Fixing question types in legacy flat format with {len(final_output['questions'])} questions")
                
                fixed_questions = []
                for question in final_output['questions']:
                    fixed_question = question.copy()
                    if 'type' in fixed_question:
                        original_type = fixed_question['type']
                        if original_type not in [t.value for t in QuestionType]:
                            mapped_type = type_mapping.get(original_type.lower(), 'unknown')
                            logger.info(f"🔧 [Document Parser] Mapped question type '{original_type}' to '{mapped_type}'")
                            
                            # If still unknown, use intelligent inference
                            if mapped_type == 'unknown':
                                inferred_type = self._infer_question_type_from_content(fixed_question)
                                logger.info(f"🔧 [Document Parser] Intelligently inferred type '{inferred_type}' for question: {fixed_question.get('text', '')[:50]}...")
                                mapped_type = inferred_type
                            
                            fixed_question['type'] = mapped_type
                    
                    # Fix scale_labels format
                    if 'scale_labels' in fixed_question:
                        scale_labels = fixed_question['scale_labels']
                        if isinstance(scale_labels, list) and len(scale_labels) >= 2:
                            fixed_question['scale_labels'] = {
                                "min": scale_labels[0],
                                "max": scale_labels[-1]
                            }
                        elif not isinstance(scale_labels, dict):
                            fixed_question.pop('scale_labels', None)
                    
                    fixed_questions.append(fixed_question)
                final_output['questions'] = fixed_questions
            
            fixed_data["final_output"] = final_output
        else:
            # Legacy format
            if 'questions' in fixed_data:
                fixed_questions = []
                for question in fixed_data['questions']:
                    fixed_question = question.copy()
                    if 'type' in fixed_question:
                        original_type = fixed_question['type']
                        if original_type not in [t.value for t in QuestionType]:
                            mapped_type = type_mapping.get(original_type.lower(), 'unknown')
                            logger.info(f"🔧 [Document Parser] Mapped question type '{original_type}' to '{mapped_type}'")
                            
                            # If still unknown, use intelligent inference
                            if mapped_type == 'unknown':
                                inferred_type = self._infer_question_type_from_content(fixed_question)
                                logger.info(f"🔧 [Document Parser] Intelligently inferred type '{inferred_type}' for question: {fixed_question.get('text', '')[:50]}...")
                                mapped_type = inferred_type
                            
                            fixed_question['type'] = mapped_type
                    
                    # Fix scale_labels format
                    if 'scale_labels' in fixed_question:
                        scale_labels = fixed_question['scale_labels']
                        if isinstance(scale_labels, list) and len(scale_labels) >= 2:
                            fixed_question['scale_labels'] = {
                                "min": scale_labels[0],
                                "max": scale_labels[-1]
                            }
                        elif not isinstance(scale_labels, dict):
                            fixed_question.pop('scale_labels', None)
                    
                    fixed_questions.append(fixed_question)
                fixed_data['questions'] = fixed_questions
        
        return fixed_data
    
    def _fix_routing_fields(self, survey_data: Dict[str, Any]) -> Dict[str, Any]:
        """Fix routing fields by converting strings to proper dict structure or setting to null."""
        logger.info(f"🔧 [Document Parser] Attempting to fix routing fields")
        
        fixed_data = survey_data.copy()
        
        # Handle both new format (with final_output) and legacy format
        if "final_output" in fixed_data:
            final_output = fixed_data["final_output"].copy()
            if 'questions' in final_output:
                fixed_questions = []
                for question in final_output['questions']:
                    fixed_question = question.copy()
                    if 'routing' in fixed_question and isinstance(fixed_question['routing'], str):
                        # Convert string routing to null (we're being less strict)
                        logger.info(f"🔧 [Document Parser] Converting string routing to null for question {fixed_question.get('id', 'unknown')}")
                        fixed_question['routing'] = None
                    fixed_questions.append(fixed_question)
                final_output['questions'] = fixed_questions
            fixed_data["final_output"] = final_output
        else:
            # Legacy format
            if 'questions' in fixed_data:
                fixed_questions = []
                for question in fixed_data['questions']:
                    fixed_question = question.copy()
                    if 'routing' in fixed_question and isinstance(fixed_question['routing'], str):
                        # Convert string routing to null (we're being less strict)
                        logger.info(f"🔧 [Document Parser] Converting string routing to null for question {fixed_question.get('id', 'unknown')}")
                        fixed_question['routing'] = None
                    fixed_questions.append(fixed_question)
                fixed_data['questions'] = fixed_questions
        
        return fixed_data
    
    def create_rfq_extraction_prompt(self, document_text: str) -> str:
        """Create the system prompt for RFQ-specific data extraction."""

        prompt = f"""You are an expert research consultant. Extract RFQ (Request for Quotation) information from the following document and convert it to structured data.

CRITICAL: You MUST return ONLY valid JSON. No explanations, no markdown, no backticks. Do NOT output character arrays or individual characters separated by spaces. Return a complete JSON string that can be parsed by json.loads().

DOCUMENT TEXT:
{document_text}

FIELD PRIORITIZATION (extract in this order of importance):
CRITICAL FIELDS (highest priority - extract first):
1. title - The main title or subject of the RFQ (usually in headers or at the beginning)
2. description - Core description of what research is needed (often the largest text block)

BUSINESS CONTEXT FIELDS (Critical):
3. company_product_background - Company background, product context, business overview
4. business_problem - What business challenge or problem needs to be solved
5. business_objective - What the business wants to achieve from this research
6. stakeholder_requirements - Key stakeholder needs and requirements
7. decision_criteria - What defines success for this research
8. budget_range - Budget constraints (under_10k, 10k_50k, 50k_100k, 100k_plus)
9. timeline_constraints - Timeline requirements (rush, standard, flexible)

RESEARCH OBJECTIVES FIELDS (High Priority):
10. research_audience - Target audience, demographics, respondent segments
11. success_criteria - Desired outcome and success criteria
12. key_research_questions - Key research questions and considerations (as array)
13. success_metrics - How research success will be measured
14. validation_requirements - What validation is needed
15. measurement_approach - Research approach (quantitative, qualitative, mixed_methods)

ADVANCED CLASSIFICATION FIELDS (High Priority):
16. industry_classification - Industry type (free text - extract from company background, product context)
17. respondent_classification - Target respondent type (free text - extract from research audience, demographics)
18. methodology_tags - Methodology tags and labels (as array)

METHODOLOGY FIELDS (High Priority):
19. primary_method - Research methodology (van_westendorp, gabor_granger, conjoint, basic_survey)
20. stimuli_details - Concept details, price ranges, stimuli information
21. methodology_requirements - Additional methodology notes and requirements
22. complexity_level - Research complexity (simple, standard, advanced)
23. required_methodologies - Specific methodologies required (as array)
24. sample_size_target - Target number of respondents

SURVEY REQUIREMENTS FIELDS (Medium Priority):
25. sample_plan - Sample structure, LOI, recruiting criteria
26. must_have_questions - Must-have questions per respondent type (as array)
27. screener_requirements - Screener and respondent qualification requirements, tagging rules, piping logic
28. completion_time_target - Target completion time (5_10_min, 10_15_min, 15_25_min, 25_plus_min)
29. device_compatibility - Device requirements (mobile_first, desktop_first, both)
30. accessibility_requirements - Accessibility needs (standard, enhanced, full_compliance)
31. data_quality_requirements - Data quality standards (basic, standard, premium)

SURVEY STRUCTURE FIELDS (Medium Priority):
32. qnr_sections_detected - Required QNR sections (as array)
33. text_requirements_detected - Required text introductions (as array)

SURVEY LOGIC FIELDS (Medium Priority):
34. requires_piping_logic - Whether piping/carry-forward logic is needed (boolean)
35. requires_sampling_logic - Whether sampling/randomization logic is needed (boolean)
36. requires_screener_logic - Whether screener/qualification logic is needed (boolean)
37. custom_logic_requirements - Custom logic requirements and specifications

BRAND USAGE FIELDS (Medium Priority):
38. brand_recall_required - Whether brand recall questions are needed (boolean)
39. brand_awareness_funnel - Whether brand awareness funnel is needed (boolean)
40. brand_product_satisfaction - Whether brand/product satisfaction is needed (boolean)
41. usage_frequency_tracking - Whether usage frequency tracking is needed (boolean)

RULES AND DEFINITIONS FIELDS (Medium Priority):
42. rules_and_definitions - Rules, definitions, jargon feed, special terms, terminology requirements

SIMPLIFIED FIELD-SPECIFIC EXTRACTION GUIDANCE:

CRITICAL FIELDS (extract with high confidence):

TITLE (Critical):
- Keywords: "title", "project", "study", "research", headers, bold text
- Patterns: "Research Study", "Market Research", "RFQ:", document title
- Strategy: Extract from document headers, subject lines, first prominent text

DESCRIPTION (Critical):
- Keywords: "overview", "description", "purpose", "objective", "we need"
- Patterns: Introduction paragraphs, executive summary text
- Strategy: Extract comprehensive overview paragraphs

COMPANY_PRODUCT_BACKGROUND (Critical):
- Keywords: "company", "background", "business", "product", "about us"
- Patterns: "Company background", "Business context", "About [company]"
- Strategy: Extract company and product context paragraphs

BUSINESS_PROBLEM (Critical):
- Keywords: "problem", "challenge", "issue", "need", "goal"
- Patterns: "Business challenge", "We need to", "The problem is"
- Strategy: Extract problem statement sentences

BUSINESS_OBJECTIVE (Critical):
- Keywords: "objective", "goal", "achieve", "outcome", "want to"
- Patterns: "Business objective", "We want to achieve", "The goal is"
- Strategy: Extract objective statements

STAKEHOLDER_REQUIREMENTS (Critical):
- Keywords: "stakeholder", "decision maker", "key users", "requirements"
- Patterns: "Stakeholder needs", "Decision maker requirements", "Key user needs"
- Strategy: Extract stakeholder and decision maker requirements

DECISION_CRITERIA (Critical):
- Keywords: "success", "criteria", "measure", "evaluate", "judge"
- Patterns: "Success criteria", "How we measure success", "Decision criteria"
- Strategy: Extract success measurement criteria

BUDGET_RANGE (Critical):
- Keywords: "budget", "cost", "price", "under", "less than", "more than"
- Patterns: "Budget: $X", "Under $10k", "Less than $50k", "Budget range"
- Strategy: Extract budget information and map to ranges

TIMELINE_CONSTRAINTS (Critical):
- Keywords: "timeline", "deadline", "urgent", "rush", "flexible", "standard"
- Patterns: "Timeline: X weeks", "Urgent deadline", "Flexible timeline"
- Strategy: Extract timeline information and map to constraints

RESEARCH_AUDIENCE (High Priority):
- Keywords: "audience", "target", "respondents", "participants", "demographics"
- Patterns: "Target audience", "Respondent profile", demographic descriptions
- Strategy: Extract audience descriptions

SUCCESS_CRITERIA (High Priority):
- Keywords: "success", "outcome", "result", "achieve", "deliverable"
- Patterns: "Success criteria", "Expected outcomes", "What we want to achieve"
- Strategy: Extract success criteria and expected outcomes

KEY_RESEARCH_QUESTIONS (High Priority):
- Keywords: "questions", "research questions", "objectives", "goals"
- Patterns: "Key questions:", "Research objectives:", numbered/bulleted lists
- Strategy: Extract research questions as array

SUCCESS_METRICS (High Priority):
- Keywords: "metrics", "measure", "KPI", "track", "monitor"
- Patterns: "Success metrics", "How we measure", "Key performance indicators"
- Strategy: Extract measurement and tracking requirements

VALIDATION_REQUIREMENTS (High Priority):
- Keywords: "validate", "verify", "confirm", "test", "check"
- Patterns: "Validation requirements", "Need to verify", "Confirm results"
- Strategy: Extract validation and verification needs

MEASUREMENT_APPROACH (High Priority):
- Keywords: "quantitative", "qualitative", "mixed", "survey", "interview"
- Patterns: "Quantitative research", "Qualitative approach", "Mixed methods"
- Strategy: Extract research approach and map to options

INDUSTRY_CLASSIFICATION (High Priority):
- Keywords: "industry", "sector", "market", "business", "company type", "vertical"
- Patterns: "Technology company", "Healthcare industry", "Financial services", "Retail sector"
- Strategy: Extract industry context from company background, product descriptions, business context
- Examples: "Technology", "Healthcare", "Financial Services", "Retail", "Automotive", "Education"

RESPONDENT_CLASSIFICATION (High Priority):
- Keywords: "respondents", "participants", "audience", "customers", "users", "consumers", "professionals"
- Patterns: "Target consumers", "B2B decision makers", "Healthcare professionals", "Students", "General public"
- Strategy: Extract respondent type from research audience, demographics, target market descriptions
- Examples: "B2C Consumers", "B2B Professionals", "Healthcare Workers", "Students", "General Public"

METHODOLOGY_TAGS (High Priority):
- Keywords: "methodology", "approach", "technique", "method", "conjoint", "van westendorp", "gabor granger", "pricing", "choice modeling"
- Patterns: "Conjoint analysis", "Van Westendorp pricing", "Gabor-Granger method", "Choice modeling", "Price sensitivity"
- Strategy: Extract methodology names and map to standardized tags
- Examples: ["conjoint", "van_westendorp", "gabor_granger", "choice_modeling", "pricing"]

PRIMARY_METHOD (High Priority):
- Keywords: "van westendorp", "gabor granger", "conjoint", "survey", "interview"
- Patterns: "Use van westendorp", "Conjoint analysis", "Price sensitivity"
- Strategy: Extract methodology and map to options

STIMULI_DETAILS (High Priority):
- Keywords: "concept", "stimuli", "product", "price", "feature"
- Patterns: "Concept details", "Product stimuli", "Price points"
- Strategy: Extract concept and stimuli information

METHODOLOGY_REQUIREMENTS (High Priority):
- Keywords: "methodology", "approach", "technique", "method", "process"
- Patterns: "Methodology requirements", "Research approach", "Technique needed"
- Strategy: Extract methodology-specific requirements

COMPLEXITY_LEVEL (High Priority):
- Keywords: "simple", "standard", "advanced", "complex", "basic"
- Patterns: "Simple research", "Advanced analysis", "Complex methodology"
- Strategy: Extract complexity level and map to options

REQUIRED_METHODOLOGIES (High Priority):
- Keywords: "methodologies", "techniques", "approaches", "methods"
- Patterns: "Required methodologies", "Use techniques", "Approaches needed"
- Strategy: Extract methodology requirements as array

SAMPLE_SIZE_TARGET (High Priority):
- Keywords: "sample", "respondents", "participants", "n=", "number"
- Patterns: "Sample size: X", "N=500", "Recruit X participants"
- Strategy: Extract target sample size

SAMPLE_PLAN (Medium Priority):
- Keywords: "sample", "LOI", "length of interview", "recruiting", "criteria"
- Patterns: "Sample size", "LOI: X minutes", "Recruit X participants"
- Strategy: Extract sampling specifications

MUST_HAVE_QUESTIONS (Medium Priority):
- Keywords: "questions", "must have", "required", "essential", "key"
- Patterns: "Must-have questions", "Required questions", "Essential questions"
- Strategy: Extract required questions as array

COMPLETION_TIME_TARGET (Medium Priority):
- Keywords: "time", "minutes", "duration", "length", "LOI"
- Patterns: "5-10 minutes", "15 minute survey", "LOI: 10 minutes"
- Strategy: Extract completion time and map to options

DEVICE_COMPATIBILITY (Medium Priority):
- Keywords: "mobile", "desktop", "device", "platform", "responsive"
- Patterns: "Mobile-first", "Desktop only", "All devices"
- Strategy: Extract device requirements and map to options

ACCESSIBILITY_REQUIREMENTS (Medium Priority):
- Keywords: "accessibility", "ADA", "compliance", "inclusive", "accessible"
- Patterns: "Accessibility requirements", "ADA compliant", "Inclusive design"
- Strategy: Extract accessibility needs and map to options

DATA_QUALITY_REQUIREMENTS (Medium Priority):
- Keywords: "quality", "data quality", "validation", "verification", "standards"
- Patterns: "Data quality standards", "Quality requirements", "Validation needed"
- Strategy: Extract data quality requirements and map to options

SCREENER_REQUIREMENTS (Medium Priority):
- Keywords: "screener", "qualification", "eligibility", "criteria", "filter", "piping", "tagging"
- Patterns: "Screener requirements", "Qualification criteria", "Eligibility rules", "Respondent tagging", "Piping logic"
- Strategy: Extract screener and qualification requirements, tagging rules, piping logic specifications

QNR_SECTIONS_DETECTED (Medium Priority):
- Keywords: "sections", "questionnaire", "survey", "screener", "demographics"
- Patterns: "Survey sections", "Questionnaire structure", "Required sections"
- Strategy: Auto-detect required QNR sections based on methodology

TEXT_REQUIREMENTS_DETECTED (Medium Priority):
- Keywords: "introduction", "instructions", "confidentiality", "study intro"
- Patterns: "Study introduction", "Confidentiality agreement", "Instructions"
- Strategy: Auto-detect required text blocks based on methodology

REQUIRES_PIPING_LOGIC (Medium Priority):
- Keywords: "pipe", "carry forward", "previous answer", "based on response"
- Patterns: "Pipe responses", "Carry forward answers", "Based on previous"
- Strategy: Detect piping logic requirements

REQUIRES_SAMPLING_LOGIC (Medium Priority):
- Keywords: "random", "quota", "balance", "representative sample"
- Patterns: "Random sampling", "Quota requirements", "Balanced sample"
- Strategy: Detect sampling logic requirements

REQUIRES_SCREENER_LOGIC (Medium Priority):
- Keywords: "qualify", "screen out", "terminate", "continue if", "skip if"
- Patterns: "Qualify respondents", "Screen out", "Terminate if"
- Strategy: Detect screener logic requirements

CUSTOM_LOGIC_REQUIREMENTS (Medium Priority):
- Keywords: "logic", "skip pattern", "routing", "conditional", "branching"
- Patterns: "Skip patterns", "Conditional logic", "Routing requirements"
- Strategy: Extract custom logic requirements

BRAND_RECALL_REQUIRED (Medium Priority):
- Keywords: "brand awareness", "top of mind", "unaided", "aided recall"
- Patterns: "Brand awareness", "Top of mind", "Unaided recall"
- Strategy: Detect brand recall requirements

BRAND_AWARENESS_FUNNEL (Medium Priority):
- Keywords: "awareness", "consideration", "trial", "purchase", "loyalty"
- Patterns: "Awareness funnel", "Consideration", "Purchase intent"
- Strategy: Detect brand awareness funnel requirements

BRAND_PRODUCT_SATISFACTION (Medium Priority):
- Keywords: "satisfaction", "rating", "experience", "recommend", "NPS"
- Patterns: "Satisfaction rating", "Product experience", "Recommendation"
- Strategy: Detect brand/product satisfaction requirements

USAGE_FREQUENCY_TRACKING (Medium Priority):
- Keywords: "how often", "frequency", "usage", "habits", "occasions"
- Patterns: "Usage frequency", "How often", "Usage habits"
- Strategy: Detect usage frequency tracking requirements

RULES_AND_DEFINITIONS (Medium Priority):
- Keywords: "rules", "definitions", "terminology", "jargon", "terms", "glossary", "specifications"
- Patterns: "Rules and definitions", "Terminology requirements", "Jargon feed", "Special terms", "Definitions needed"
- Strategy: Extract rules, definitions, terminology requirements, jargon specifications, special terms that need definition

SURVEY STRUCTURE FIELDS (methodology-driven):

QNR_SECTIONS_DETECTED (Medium):
- Keywords: "questionnaire", "survey sections", "screener", "demographics", "concept"
- Patterns: Based on detected methodologies (see methodology mapping below)
- Strategy: Auto-detect required QNR sections based on methodology
- Default sections: ["sample_plan", "screener", "brand_awareness", "concept_exposure", "methodology_section", "additional_questions", "programmer_instructions"]

TEXT_REQUIREMENTS_DETECTED (Medium):
- Keywords: "introduction", "instructions", "confidentiality", "study intro", "concept presentation"
- Patterns: Based on methodology requirements (see text requirements mapping below)
- Strategy: Auto-detect required text blocks based on methodology

SURVEY_LOGIC_REQUIREMENTS (Medium):
- requires_piping_logic: Keywords: "pipe", "carry forward", "previous answer", "based on response"
- requires_sampling_logic: Keywords: "random", "quota", "balance", "representative sample"
- requires_screener_logic: Keywords: "qualify", "screen out", "terminate", "continue if", "skip if"
- custom_logic_requirements: Keywords: "logic", "skip pattern", "routing", "conditional"
- Strategy: Infer logic needs from content complexity and methodology

BRAND_USAGE_REQUIREMENTS (Medium):
- brand_recall_required: Keywords: "brand awareness", "top of mind", "unaided", "aided recall"
- brand_awareness_funnel: Keywords: "awareness", "consideration", "trial", "purchase", "loyalty"
- brand_product_satisfaction: Keywords: "satisfaction", "rating", "experience", "recommend"
- usage_frequency_tracking: Keywords: "how often", "frequency", "usage", "habits", "occasions"
- Strategy: Detect brand/product research focus

ENHANCED BUSINESS CONTEXT (Medium):
- stakeholder_requirements: Keywords: "stakeholders", "requirements", "needs", "expectations"
- decision_criteria: Keywords: "criteria", "decision", "evaluate", "success metrics"
- budget_range: Keywords: "budget", "cost", "investment", "under", "over", "$", "k", "million"
- timeline_constraints: Keywords: "urgent", "deadline", "timeline", "rush", "flexible", "standard"

ENHANCED RESEARCH OBJECTIVES (Medium):
- success_metrics: Keywords: "success", "metrics", "KPI", "measure", "target", "goal"
- validation_requirements: Keywords: "validate", "verify", "confirm", "test", "proof"
- measurement_approach: Keywords: "quantitative", "qualitative", "mixed", "approach", "method"

ENHANCED METHODOLOGY (Medium):
- complexity_level: Keywords: "simple", "standard", "complex", "advanced", "basic", "sophisticated"
- required_methodologies: Keywords: Multiple methodology names in array format
- sample_size_target: Keywords: "sample", "respondents", "participants", "n=", "size"

ENHANCED SURVEY REQUIREMENTS (Medium):
- completion_time_target: Keywords: "time", "minutes", "duration", "length", "LOI"
- device_compatibility: Keywords: "mobile", "desktop", "tablet", "device", "responsive"
- accessibility_requirements: Keywords: "accessibility", "ADA", "508", "WCAG", "accessible"
- data_quality_requirements: Keywords: "quality", "validation", "attention", "checks"

EXTRACTION STRATEGY:
1. First scan for CRITICAL FIELDS using field-specific patterns above
2. Look for explicit section headers that match field types
3. Identify numbered/bulleted lists that may contain objectives or constraints
4. Use field-specific language patterns for targeted extraction
5. For missing CRITICAL FIELDS, be more aggressive - synthesize from available context
6. For MEDIUM PRIORITY fields, extract if clear evidence exists
7. Apply METHODOLOGY-BASED AUTO-DETECTION for survey structure fields using mappings below

METHODOLOGY-TO-QNR-SECTIONS MAPPING:
- concept_test: ["sample_plan", "screener", "brand_awareness", "concept_exposure", "methodology_section", "additional_questions", "programmer_instructions"]
- conjoint: ["sample_plan", "screener", "concept_exposure", "methodology_section", "additional_questions", "programmer_instructions"]
- van_westendorp/gabor_granger: ["sample_plan", "screener", "brand_awareness", "methodology_section", "additional_questions", "programmer_instructions"]
- brand_tracker: ["sample_plan", "screener", "brand_awareness", "additional_questions", "programmer_instructions"]
- max_diff: ["sample_plan", "screener", "methodology_section", "additional_questions", "programmer_instructions"]

METHODOLOGY-TO-TEXT-REQUIREMENTS MAPPING:
- concept_test: ["Study_Intro", "Concept_Intro"]
- conjoint: ["Study_Intro", "Confidentiality_Agreement"]
- van_westendorp/gabor_granger: ["Study_Intro", "Product_Usage"]
- brand_tracker: ["Study_Intro", "Product_Usage"]
- max_diff: ["Study_Intro"]
- All methodologies: Always include "Study_Intro" (mandatory)

SURVEY_LOGIC AUTO-DETECTION RULES:
- Requires Piping Logic: If conjoint, complex branching, or "based on previous response" detected
- Requires Sampling Logic: If quota requirements, randomization, or balanced design mentioned
- Requires Screener Logic: If qualification criteria, screening questions, or termination logic present

BRAND_USAGE AUTO-DETECTION RULES:
- Brand Recall Required: If brand awareness, recall, or brand tracking mentioned
- Brand Awareness Funnel: If awareness-to-purchase funnel or brand consideration mentioned
- Brand Product Satisfaction: If satisfaction, ratings, or experience evaluation mentioned
- Usage Frequency Tracking: If usage patterns, frequency, or habits research mentioned

CONFIDENCE SCORING GUIDELINES:
- 0.9-1.0: Explicit mention with clear mapping (exact field labels or unambiguous content)
- 0.7-0.8: Strong contextual evidence (business language patterns, section positioning)
- 0.5-0.6: Reasonable inference from surrounding context
- 0.3-0.4: Weak inference or partial information
- 0.0-0.2: Speculative or very uncertain

INSTRUCTIONS:
1. Extract the following RFQ components following the prioritization above:
   - Research objectives and goals (CRITICAL)
   - Business context and background (HIGH)
   - Target audience and demographics (HIGH)
   - Constraints (budget, timeline, sample size, methodology) (MEDIUM)
   - Stakeholder requirements
   - Success metrics
   - Preferred or excluded methodologies

2. Structure the output as JSON with the following schema:
   - confidence: number (0.0-1.0) - how confident you are in the extraction
   - identified_sections: object with potential matches for each RFQ component
   - extracted_entities: object with lists of identified entities (stakeholders, industries, etc.)
   - field_mappings: array of specific field mappings with confidence scores

3. For each field mapping, include:
   - field: the RFQ field name (prioritized as above)
   - value: the extracted value
   - confidence: extraction confidence (0.0-1.0)
   - source: the text snippet this was extracted from
   - reasoning: why this text maps to this field
   - priority: field priority level (critical/high/medium)

UNMAPPED_CONTEXT FIELD:
- Purpose: Capture any useful information that doesn't fit into structured fields but could help with survey generation
- Examples: Methodology preferences, tone requirements, special constraints, stakeholder notes, cultural considerations, brand guidelines, specific terminology preferences
- Strategy: Look for contextual details, preferences, or requirements that would influence survey design but aren't captured in the structured fields
- Limit: Maximum 200 words, focus on actionable information for survey generation
- If no additional context exists, use empty string ""

EXPECTED JSON STRUCTURE:
{{
  "confidence": 0.85,
  "identified_sections": {{
    "objectives": {{
      "confidence": 0.9,
      "source_text": "extracted text snippet",
      "source_section": "section where found",
      "extracted_data": ["objective 1", "objective 2"]
    }},
    "business_context": {{ ... }},
    "target_audience": {{ ... }},
    "constraints": {{ ... }},
    "methodologies": {{ ... }}
  }},
  "extracted_entities": {{
    "stakeholders": ["decision maker", "end user"],
    "industries": ["technology", "healthcare"],
    "research_types": ["market research", "pricing study"],
    "methodologies": ["van westendorp", "conjoint analysis"],
    "industry_classification": "Technology",
    "respondent_classification": "B2C Consumers"
  }},
  "field_mappings": [
    {{
      "field": "title",
      "value": "Product Feature Research Study",
      "confidence": 0.95,
      "source": "Title: Product Feature Research Study",
      "reasoning": "Clear title in document header",
      "priority": "critical"
    }},
    {{
      "field": "description",
      "value": "Research to understand customer preferences for new product features",
      "confidence": 0.85,
      "source": "We need to understand what features customers value most...",
      "reasoning": "Main research description in introduction",
      "priority": "critical"
    }},
    {{
      "field": "company_product_background",
      "value": "TechCorp is a leading software company developing productivity tools for small businesses.",
      "confidence": 0.80,
      "source": "About TechCorp: We are a leading software company...",
      "reasoning": "Company background section with clear business context",
      "priority": "critical"
    }},
    {{
      "field": "business_problem",
      "value": "Need to prioritize which features to develop for next product release",
      "confidence": 0.85,
      "source": "The challenge we face is determining which features...",
      "reasoning": "Clear problem statement in business context",
      "priority": "critical"
    }},
    {{
      "field": "business_objective",
      "value": "Determine which features to prioritize for next product release",
      "confidence": 0.85,
      "source": "Our objective is to determine feature priorities...",
      "reasoning": "Clear business objective statement",
      "priority": "critical"
    }},
    {{
      "field": "stakeholder_requirements",
      "value": "Product managers, engineering team, and executive leadership need input",
      "confidence": 0.75,
      "source": "Stakeholders include product managers and engineering...",
      "reasoning": "Stakeholder requirements mentioned in context",
      "priority": "critical"
    }},
    {{
      "field": "decision_criteria",
      "value": "Feature prioritization based on customer value and development effort",
      "confidence": 0.80,
      "source": "We need to balance customer value with development effort...",
      "reasoning": "Decision criteria implied from business context",
      "priority": "critical"
    }},
    {{
      "field": "budget_range",
      "value": "10k_50k",
      "confidence": 0.70,
      "source": "Budget: $25,000 for this research project",
      "reasoning": "Budget mentioned falls within 10k-50k range",
      "priority": "critical"
    }},
    {{
      "field": "timeline_constraints",
      "value": "standard",
      "confidence": 0.75,
      "source": "Timeline: 6 weeks for completion",
      "reasoning": "Standard timeline mentioned, not urgent or flexible",
      "priority": "critical"
    }},
    {{
      "field": "research_audience",
      "value": "Small business owners, 25-50 years old, currently using productivity software",
      "confidence": 0.90,
      "source": "Target participants: small business owners aged 25-50...",
      "reasoning": "Clear demographic and behavioral criteria",
      "priority": "high"
    }},
    {{
      "field": "success_criteria",
      "value": "Clear feature prioritization ranking with confidence scores",
      "confidence": 0.85,
      "source": "Success will be measured by clear feature ranking...",
      "reasoning": "Success criteria explicitly mentioned",
      "priority": "high"
    }},
    {{
      "field": "key_research_questions",
      "value": ["Which features are most important to customers?", "What is the relative value of each feature?", "How do features interact with each other?"],
      "confidence": 0.90,
      "source": "Key questions: 1) Feature importance 2) Relative value 3) Feature interactions",
      "reasoning": "Research questions clearly listed",
      "priority": "high"
    }},
    {{
      "field": "success_metrics",
      "value": "Feature importance scores, relative value rankings, and interaction effects",
      "confidence": 0.80,
      "source": "Metrics: importance scores, rankings, interaction effects",
      "reasoning": "Success metrics mentioned in context",
      "priority": "high"
    }},
    {{
      "field": "validation_requirements",
      "value": "Statistical significance testing and confidence intervals",
      "confidence": 0.75,
      "source": "Need statistical validation with confidence intervals",
      "reasoning": "Validation requirements mentioned",
      "priority": "high"
    }},
    {{
      "field": "measurement_approach",
      "value": "quantitative",
      "confidence": 0.85,
      "source": "Quantitative research approach using conjoint analysis",
      "reasoning": "Quantitative approach explicitly mentioned",
      "priority": "high"
    }},
    {{
      "field": "primary_method",
      "value": "conjoint",
      "confidence": 0.90,
      "source": "We want to use conjoint analysis to understand feature trade-offs",
      "reasoning": "Explicit mention of conjoint analysis methodology",
      "priority": "high"
    }},
    {{
      "field": "stimuli_details",
      "value": "Product features: AI automation, mobile app, cloud sync, advanced analytics",
      "confidence": 0.85,
      "source": "Features to test: AI automation, mobile app, cloud sync, analytics",
      "reasoning": "Product features clearly listed for conjoint stimuli",
      "priority": "high"
    }},
    {{
      "field": "methodology_requirements",
      "value": "Conjoint analysis with 16 choice tasks, randomized presentation",
      "confidence": 0.80,
      "source": "16 choice tasks with randomized presentation order",
      "reasoning": "Methodology requirements specified",
      "priority": "high"
    }},
    {{
      "field": "complexity_level",
      "value": "advanced",
      "confidence": 0.75,
      "source": "Complex conjoint analysis with multiple attributes",
      "reasoning": "Advanced complexity due to conjoint methodology",
      "priority": "high"
    }},
    {{
      "field": "required_methodologies",
      "value": ["conjoint", "choice_modeling", "statistical_analysis"],
      "confidence": 0.85,
      "source": "Conjoint analysis, choice modeling, and statistical analysis required",
      "reasoning": "Required methodologies explicitly mentioned",
      "priority": "high"
    }},
    {{
      "field": "sample_size_target",
      "value": "500",
      "confidence": 0.90,
      "source": "Sample size: 500 respondents",
      "reasoning": "Sample size explicitly stated",
      "priority": "high"
    }},
    {{
      "field": "sample_plan",
      "value": "500 small business owners, 15-minute LOI, recruited via panel",
      "confidence": 0.85,
      "source": "Recruit 500 small business owners, 15-minute survey via panel",
      "reasoning": "Sample plan details provided",
      "priority": "medium"
    }},
    {{
      "field": "must_have_questions",
      "value": ["Feature importance ranking", "Purchase intent", "Demographics"],
      "confidence": 0.80,
      "source": "Must include: feature ranking, purchase intent, demographics",
      "reasoning": "Required questions mentioned",
      "priority": "medium"
    }},
    {{
      "field": "completion_time_target",
      "value": "15_25_min",
      "confidence": 0.85,
      "source": "15-minute survey length",
      "reasoning": "Survey length specified as 15 minutes",
      "priority": "medium"
    }},
    {{
      "field": "device_compatibility",
      "value": "both",
      "confidence": 0.80,
      "source": "Survey should work on mobile and desktop",
      "reasoning": "Device compatibility requirements mentioned",
      "priority": "medium"
    }},
    {{
      "field": "accessibility_requirements",
      "value": "standard",
      "confidence": 0.75,
      "source": "Standard accessibility compliance needed",
      "reasoning": "Accessibility requirements mentioned",
      "priority": "medium"
    }},
    {{
      "field": "data_quality_requirements",
      "value": "premium",
      "confidence": 0.80,
      "source": "High-quality data with validation checks required",
      "reasoning": "Premium data quality requirements mentioned",
      "priority": "medium"
    }},
    {{
      "field": "qnr_sections_detected",
      "value": ["sample_plan", "screener", "concept_exposure", "methodology_section", "additional_questions", "programmer_instructions"],
      "confidence": 0.85,
      "source": "Based on conjoint methodology requirements",
      "reasoning": "Auto-detected based on conjoint analysis methodology",
      "priority": "medium"
    }},
    {{
      "field": "text_requirements_detected",
      "value": ["Study_Intro", "Confidentiality_Agreement"],
      "confidence": 0.80,
      "source": "Standard conjoint study requirements",
      "reasoning": "Auto-detected based on conjoint methodology",
      "priority": "medium"
    }},
    {{
      "field": "requires_piping_logic",
      "value": true,
      "confidence": 0.75,
      "source": "Conjoint analysis requires complex branching",
      "reasoning": "Conjoint methodology inherently requires piping logic",
      "priority": "medium"
    }},
    {{
      "field": "requires_sampling_logic",
      "value": false,
      "confidence": 0.80,
      "source": "No specific sampling requirements mentioned",
      "reasoning": "No quota or randomization requirements detected",
      "priority": "medium"
    }},
    {{
      "field": "requires_screener_logic",
      "value": true,
      "confidence": 0.85,
      "source": "Need to screen for small business owners using productivity software",
      "reasoning": "Screener logic needed for target audience qualification",
      "priority": "medium"
    }},
    {{
      "field": "custom_logic_requirements",
      "value": "Randomize choice task order, balance attribute levels",
      "confidence": 0.80,
      "source": "Randomize choice tasks and balance attribute levels",
      "reasoning": "Custom logic requirements for conjoint analysis",
      "priority": "medium"
    }},
    {{
      "field": "brand_recall_required",
      "value": false,
      "confidence": 0.80,
      "source": "Focus on product features, not brand awareness",
      "reasoning": "Feature research doesn't require brand recall questions",
      "priority": "medium"
    }},
    {{
      "field": "brand_awareness_funnel",
      "value": false,
      "confidence": 0.80,
      "source": "No brand awareness measurement needed",
      "reasoning": "Feature research doesn't require brand awareness funnel",
      "priority": "medium"
    }},
    {{
      "field": "brand_product_satisfaction",
      "value": false,
      "confidence": 0.80,
      "source": "Focus on feature preferences, not satisfaction",
      "reasoning": "Feature research doesn't require satisfaction measurement",
      "priority": "medium"
    }},
    {{
      "field": "usage_frequency_tracking",
      "value": false,
      "confidence": 0.80,
      "source": "No usage frequency tracking needed",
      "reasoning": "Feature research doesn't require usage frequency tracking",
      "priority": "medium"
    }},
    {{
      "field": "methodology_tags",
      "value": ["conjoint", "choice_modeling"],
      "confidence": 0.85,
      "source": "Conjoint analysis methodology mentioned in document",
      "reasoning": "Conjoint analysis and choice modeling methodologies detected",
      "priority": "high"
    }}
  ]
}}

IMPORTANT:
- Return ONLY valid JSON, no explanations or additional text
- Use null for missing or unclear information
- Be conservative with confidence scores
- Include source text snippets for traceability
- CRITICAL: Replace all newline characters (\n) in field values with spaces to ensure clean JSON
- Clean field values by removing excessive whitespace and line breaks

COMPLETE JSON EXAMPLE:
{{
  "confidence": 0.85,
  "field_mappings": [
    {{
      "field": "title",
      "value": "Quantitative Market Research for Contact Lens Pricing",
      "confidence": 0.95,
      "source": "Request for Quotation (RFQ): Quantitative Market Research for Contact Lens Pricing",
      "reasoning": "Clear title in document header",
      "priority": "critical"
    }},
    {{
      "field": "description", 
      "value": "Comprehensive quantitative pricing study for AIR OPTIX plus HydraGlyde contact lens product to determine optimal price points for market acceptance",
      "confidence": 0.90,
      "source": "Alcon Consumer Insights team is seeking proposals from qualified market research agencies to conduct a comprehensive quantitative pricing study",
      "reasoning": "Main research description in executive summary",
      "priority": "critical"
    }},
    {{
      "field": "research_audience",
      "value": "Current contact lens wearers (70% of sample), age 18-45 with focus on 25-34, 60% female/40% male, middle to upper-middle income, monthly replacement users seeking premium comfort",
      "confidence": 0.95,
      "source": "Primary Target Demographics: Current Contact Lens Wearers (70% of sample) Age: 18-45 years (primary focus on 25-34 age group representing 39% of wearers)",
      "reasoning": "Explicit audience details listed",
      "priority": "high"
    }},
    {{
      "field": "primary_method",
      "value": "gabor_granger",
      "confidence": 0.62,
      "source": "Assess consumer willingness to pay at various price levels and evaluating price elasticity",
      "reasoning": "WTP at various price levels and elasticity align with Gabor-Granger methodology",
      "priority": "high"
    }},
    {{
      "field": "sample_size_target",
      "value": "800-1200",
      "confidence": 0.97,
      "source": "Total Sample Size: 800-1,200 respondents",
      "reasoning": "Explicit sample size range provided",
      "priority": "high"
    }},
    {{
      "field": "screener_requirements",
      "value": "Qualified respondents must be current contact lens wearers, age 18-45, with monthly replacement usage. Include piping logic for brand awareness questions.",
      "confidence": 0.85,
      "source": "Screener requirements: Current contact lens wearers only, age 18-45, monthly replacement users",
      "reasoning": "Explicit screener criteria and piping logic requirements mentioned",
      "priority": "medium"
    }},
    {{
      "field": "rules_and_definitions",
      "value": "Define 'premium comfort' as enhanced moisture retention and all-day comfort. 'Price sensitivity' refers to willingness to pay premium prices for superior comfort features.",
      "confidence": 0.80,
      "source": "Key terms requiring definition: premium comfort, price sensitivity, enhanced features",
      "reasoning": "Specialized terminology requiring clear definitions for survey consistency",
      "priority": "medium"
    }}
  ],
  "unmapped_context": "Any additional useful information from the document that doesn't fit into the structured fields above but could be valuable for survey generation. This might include methodology preferences, tone requirements, special constraints, stakeholder notes, or other contextual details. Keep this concise (maximum 200 words) and focus on information that would help create better surveys."
}}

REMEMBER: Return ONLY the JSON structure above. No other text, explanations, or formatting.
"""
        return prompt

    async def extract_rfq_data(self, document_text: str) -> Dict[str, Any]:
        """Extract RFQ-specific data from document text using LLM."""
        logger.info(f"🎯 [Document Parser] Starting RFQ data extraction")
        try:
            prompt = self.create_rfq_extraction_prompt(document_text)
            logger.info(f"📝 [Document Parser] Created RFQ extraction prompt, length: {len(prompt)} chars")

            # Create audit context for this LLM interaction
            interaction_id = f"rfq_extraction_{uuid.uuid4().hex[:8]}"
            audit_service = LLMAuditService(self.db_session) if self.db_session else None
            
            # Try to use audit service, but don't let audit failures break the core functionality
            try:
                if audit_service:
                    async with LLMAuditContext(
                        audit_service=audit_service,
                        interaction_id=interaction_id,
                        model_name=self.model,
                        model_provider="replicate",
                        purpose="document_parsing",
                        input_prompt=prompt,
                        sub_purpose="rfq_extraction",
                        context_type="document",
                        hyperparameters={
                            "temperature": 0.1,
                            "max_tokens": 4000
                        },
                        metadata={
                            "document_length": len(document_text),
                            "prompt_length": len(prompt)
                        },
                        tags=["document_parsing", "rfq_extraction"]
                    ) as audit_context:
                        logger.info(f"🚀 [Document Parser] Calling Replicate API for RFQ extraction with auditing")
                        logger.info(f"🎯 [Document Parser] Using model: {self.model}")
                        logger.info(f"🎯 [Document Parser] Model source: DocumentParser.extract_rfq_data (audited)")
                        start_time = time.time()
                        output = await self.replicate_client.async_run(
                            self.model,
                            input={
                                "prompt": prompt,
                                "temperature": 0.1,
                                "max_tokens": 4000,
                                "system_prompt": "You are an expert at extracting structured information from research documents. You MUST return ONLY valid JSON that can be parsed by json.loads(). No explanations, no markdown, no backticks. Do NOT output character arrays or individual characters separated by spaces. The response must be a complete, valid JSON object matching the provided schema exactly."
                            }
                        )
                        
                        # Process the output and set audit context
                        response_time_ms = int((time.time() - start_time) * 1000)
                        
                        # Capture raw response immediately (unprocessed)
                        if hasattr(output, '__iter__') and not isinstance(output, str):
                            raw_response = "".join(str(chunk) for chunk in output)
                        else:
                            raw_response = str(output)
                        
                        # Process the output for further use
                        processed_output = raw_response.strip()
                        
                        # Set raw response (unprocessed) and processed output
                        audit_context.set_raw_response(raw_response)
                        audit_context.set_output(
                            output_content=processed_output
                        )
                else:
                    # Fallback without auditing
                    logger.info(f"🚀 [Document Parser] Calling Replicate API for RFQ extraction without auditing")
                    logger.info(f"🎯 [Document Parser] Using model: {self.model}")
                    logger.info(f"🎯 [Document Parser] Model source: DocumentParser.extract_rfq_data (non-audited)")
                    output = await self.replicate_client.async_run(
                        self.model,
                        input={
                            "prompt": prompt,
                            "temperature": 0.1,
                            "max_tokens": 4000,
                            "system_prompt": "You are an expert at extracting structured information from research documents. You MUST return ONLY valid JSON that can be parsed by json.loads(). No explanations, no markdown, no backticks. Do NOT output character arrays or individual characters separated by spaces. The response must be a complete, valid JSON object matching the provided schema exactly."
                        }
                    )
            except Exception as audit_error:
                # If audit fails, log the error but continue with core functionality
                logger.warning(f"⚠️ [Document Parser] Audit system failed, continuing without audit: {str(audit_error)}")
                logger.info(f"🚀 [Document Parser] Calling Replicate API for RFQ extraction without auditing (audit failed)")
                output = await self.replicate_client.async_run(
                    self.model,
                    input={
                        "prompt": prompt,
                        **get_json_optimized_hyperparameters("rfq_parsing"),
                        "system_prompt": "You are an expert at extracting structured information from research documents. You MUST return ONLY valid JSON that can be parsed by json.loads(). No explanations, no markdown, no backticks. Do NOT output character arrays or individual characters separated by spaces. The response must be a complete, valid JSON object matching the provided schema exactly."
                    }
                )

            # Process the response - use raw response if available from audit context
            if hasattr(output, '__iter__') and not isinstance(output, str):
                json_content = "".join(str(chunk) for chunk in output).strip()
            else:
                json_content = str(output).strip()
            
            # If we have a raw response from audit context, use that for JSON extraction
            # This ensures we're extracting from the unprocessed LLM output
            if 'audit_context' in locals() and hasattr(audit_context, 'raw_response') and audit_context.raw_response:
                logger.info(f"🔍 [Document Parser] Using raw response from audit context for JSON extraction")
                json_content = audit_context.raw_response
            
            # CRITICAL FIX: Handle character array output from LLM
            # Sometimes the LLM returns a character array instead of a JSON string
            if json_content.startswith('[') and json_content.endswith(']'):
                try:
                    # Try to parse as a character array and join it
                    import ast
                    char_array = ast.literal_eval(json_content)
                    if isinstance(char_array, list) and all(isinstance(c, str) for c in char_array):
                        json_content = ''.join(char_array).strip()
                        logger.info(f"🔧 [Document Parser] Converted character array to string, length: {len(json_content)}")
                        logger.info(f"🔧 [Document Parser] First 200 chars: {json_content[:200]}")
                    else:
                        logger.warning(f"⚠️ [Document Parser] Character array contains non-string elements")
                except Exception as e:
                    logger.warning(f"⚠️ [Document Parser] Failed to parse character array: {e}")
                    logger.warning(f"⚠️ [Document Parser] Raw content: {json_content[:100]}")

            logger.info(f"✅ [Document Parser] RFQ extraction response received, length: {len(json_content)} chars")
            logger.info(f"🔍 [Document Parser] Raw LLM response: {json_content[:500]}...")

            # Parse and validate JSON using robust extraction like survey generation
            rfq_data = self._extract_rfq_json(json_content)
            
            # CRITICAL FIX: Clean up newline characters in field values
            if 'field_mappings' in rfq_data:
                for mapping in rfq_data['field_mappings']:
                    if 'value' in mapping and isinstance(mapping['value'], str):
                        # Replace newlines with spaces and clean up whitespace
                        mapping['value'] = ' '.join(mapping['value'].replace('\n', ' ').split())
                        # Also clean source and reasoning fields
                        if 'source' in mapping and isinstance(mapping['source'], str):
                            mapping['source'] = ' '.join(mapping['source'].replace('\n', ' ').split())
                        if 'reasoning' in mapping and isinstance(mapping['reasoning'], str):
                            mapping['reasoning'] = ' '.join(mapping['reasoning'].replace('\n', ' ').split())
            
            # Log the cleaned JSON that will be used by the frontend
            logger.info(f"✅ [Document Parser] RFQ data parsing successful")
            logger.info(f"🔍 [Document Parser] Parsed data keys: {list(rfq_data.keys())}")
            if 'field_mappings' in rfq_data:
                logger.info(f"📊 [Document Parser] Field mappings found: {len(rfq_data['field_mappings'])}")
                for i, mapping in enumerate(rfq_data['field_mappings'][:3]):
                    field_value = mapping.get('value', '')
                    value_preview = field_value[:50] if field_value is not None else '<null>'
                    logger.info(f"  {i+1}. {mapping.get('field', 'unknown')}: {value_preview}...")
            
            # Log the complete cleaned JSON structure for frontend debugging
            logger.info(f"🎯 [Document Parser] CLEANED JSON FOR FRONTEND:")
            logger.info(f"📄 [Document Parser] Complete RFQ data structure:")
            logger.info(f"   - Confidence: {rfq_data.get('confidence', 'N/A')}")
            logger.info(f"   - Field mappings count: {len(rfq_data.get('field_mappings', []))}")
            
            # Log each field mapping in detail
            if 'field_mappings' in rfq_data:
                for i, mapping in enumerate(rfq_data['field_mappings']):
                    field_name = mapping.get('field', 'unknown')
                    field_value = mapping.get('value', '')
                    field_confidence = mapping.get('confidence', 'N/A')
                    field_priority = mapping.get('priority', 'N/A')
                    logger.info(f"   📋 Field {i+1}: {field_name}")
                    
                    # Safe handling of field_value which might be None
                    if field_value is not None:
                        value_str = str(field_value)
                        logger.info(f"      Value: {value_str[:100]}{'...' if len(value_str) > 100 else ''}")
                    else:
                        logger.info(f"      Value: <null>")
                    
                    logger.info(f"      Confidence: {field_confidence}, Priority: {field_priority}")
            
            # Log the complete JSON as a string for debugging
            try:
                import json
                json_str = json.dumps(rfq_data, indent=2, ensure_ascii=False)
                logger.info(f"🔍 [Document Parser] Complete JSON structure (first 1000 chars):")
                logger.info(f"{json_str[:1000]}{'...' if len(json_str) > 1000 else ''}")
            except Exception as json_error:
                logger.warning(f"⚠️ [Document Parser] Could not serialize JSON for logging: {json_error}")
            
            return rfq_data

        except Exception as e:
            logger.error(f"❌ [Document Parser] Failed to extract RFQ data: {str(e)}", exc_info=True)
            return self._get_fallback_rfq_structure(str(e))

    def _extract_rfq_json(self, raw_text: str) -> Dict[str, Any]:
        """
        Extract RFQ JSON from raw LLM output using multiple strategies.
        Based on the robust extraction method from survey generation.
        """
        logger.info(f"🔍 [Document Parser] Starting JSON extraction from raw text (length: {len(raw_text)})")

        # Strategy 1: Try to parse JSON directly without any sanitization
        logger.info("🔧 [Document Parser] Trying direct JSON parsing (no sanitization)...")
        try:
            result = json.loads(raw_text)
            logger.info(f"✅ [Document Parser] Direct JSON parsing succeeded! Keys: {list(result.keys())}")
            
            # Check if this is a nested structure with json_output field
            if isinstance(result, dict) and 'json_output' in result:
                logger.info("🔧 [Document Parser] Found nested json_output structure, extracting inner data...")
                inner_data = result['json_output']
                if isinstance(inner_data, dict):
                    logger.info(f"✅ [Document Parser] Successfully extracted nested RFQ data! Keys: {list(inner_data.keys())}")
                    self._validate_and_fix_rfq_structure(inner_data)
                    return inner_data
            
            self._validate_and_fix_rfq_structure(result)
            return result
        except json.JSONDecodeError as e:
            logger.info(f"⚠️ [Document Parser] Direct JSON parsing failed: {e}")
            logger.info(f"🔍 [Document Parser] JSON error at position {e.pos}: {raw_text[max(0, e.pos-50):e.pos+50]}")

        # Strategy 2: Try to extract JSON from markdown code blocks
        logger.info("🔧 [Document Parser] Trying JSON extraction from markdown...")
        try:
            import re
            # Look for JSON in markdown code blocks (more flexible pattern)
            json_match = re.search(r'```(?:json)?\s*(\{.*?\})\s*```', raw_text, re.DOTALL)
            if json_match:
                extracted_json = json_match.group(1)
                logger.info(f"🔧 [Document Parser] Found JSON in markdown code block, length: {len(extracted_json)}")
                result = json.loads(extracted_json)
                logger.info(f"✅ [Document Parser] Markdown JSON extraction succeeded!")
                self._validate_and_fix_rfq_structure(result)
                return result
        except (json.JSONDecodeError, AttributeError) as e:
            logger.info(f"⚠️ [Document Parser] Markdown JSON extraction failed: {e}")

        # Strategy 2.5: Try to extract JSON from any code block (not just markdown)
        logger.info("🔧 [Document Parser] Trying JSON extraction from any code block...")
        try:
            import re
            # Look for JSON in any code block format
            json_match = re.search(r'```\s*(\{.*?\})\s*```', raw_text, re.DOTALL)
            if json_match:
                extracted_json = json_match.group(1)
                logger.info(f"🔧 [Document Parser] Found JSON in code block, length: {len(extracted_json)}")
                result = json.loads(extracted_json)
                logger.info(f"✅ [Document Parser] Code block JSON extraction succeeded!")
                self._validate_and_fix_rfq_structure(result)
                return result
        except (json.JSONDecodeError, AttributeError) as e:
            logger.info(f"⚠️ [Document Parser] Code block JSON extraction failed: {e}")

        # Strategy 3: Try to find JSON between first { and last }
        logger.info("🔧 [Document Parser] Trying JSON extraction between braces...")
        try:
            import re
            start = raw_text.find('{')
            end = raw_text.rfind('}') + 1
            if start != -1 and end != 0 and end > start:
                extracted_json = raw_text[start:end]
                logger.info(f"🔧 [Document Parser] Found JSON between braces, length: {len(extracted_json)}")
                result = json.loads(extracted_json)
                logger.info(f"✅ [Document Parser] Braces JSON extraction succeeded!")
                self._validate_and_fix_rfq_structure(result)
                return result
        except json.JSONDecodeError as e:
            logger.info(f"⚠️ [Document Parser] Braces JSON extraction failed: {e}")

        # Strategy 3.5: Try to find JSON with more flexible pattern (handles nested braces better)
        logger.info("🔧 [Document Parser] Trying flexible JSON extraction...")
        try:
            import re
            # Look for JSON object pattern with better handling of nested braces
            json_pattern = r'\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\}'
            json_matches = re.findall(json_pattern, raw_text, re.DOTALL)
            for match in json_matches:
                if len(match) > 50:  # Reasonable minimum length
                    try:
                        result = json.loads(match)
                        logger.info(f"✅ [Document Parser] Flexible JSON extraction succeeded!")
                        self._validate_and_fix_rfq_structure(result)
                        return result
                    except json.JSONDecodeError:
                        continue
        except Exception as e:
            logger.info(f"⚠️ [Document Parser] Flexible JSON extraction failed: {e}")

        # Strategy 4: Try gentle sanitization and parse
        logger.info("🔧 [Document Parser] Trying gentle sanitization...")
        try:
            sanitized = self._gentle_sanitize_json(raw_text)
            result = json.loads(sanitized)
            logger.info(f"✅ [Document Parser] Sanitized JSON parsing succeeded!")
            self._validate_and_fix_rfq_structure(result)
            return result
        except json.JSONDecodeError as e:
            logger.info(f"⚠️ [Document Parser] Sanitized JSON parsing failed: {e}")

        # Strategy 5: Try to extract field mappings array specifically
        logger.info("🔧 [Document Parser] Trying field mappings extraction...")
        try:
            field_mappings = self._extract_field_mappings_from_text(raw_text)
            if field_mappings:
                logger.info(f"✅ [Document Parser] Extracted {len(field_mappings)} field mappings from text")
                return {
                    "confidence": 0.9,
                    "identified_sections": {},
                    "extracted_entities": {
                        "stakeholders": [],
                        "industries": [],
                        "research_types": [],
                        "methodologies": []
                    },
                    "field_mappings": field_mappings
                }
        except Exception as e:
            logger.info(f"⚠️ [Document Parser] Field mappings extraction failed: {e}")

        # All strategies failed - return fallback
        logger.error(f"❌ [Document Parser] All JSON extraction strategies failed!")
        logger.error(f"❌ [Document Parser] Raw response: {raw_text[:1000]}...")
        return self._get_fallback_rfq_structure("All JSON extraction strategies failed")

    def _gentle_sanitize_json(self, raw_text: str) -> str:
        """Gentle JSON sanitization to fix common LLM output issues"""
        import re

        logger.info(f"🧹 [Document Parser] Starting gentle sanitization of text (length: {len(raw_text)})")

        # Remove any text before the first {
        start_idx = raw_text.find('{')
        if start_idx > 0:
            raw_text = raw_text[start_idx:]
            logger.info(f"🧹 [Document Parser] Removed text before first brace")

        # Remove any text after the last }
        end_idx = raw_text.rfind('}')
        if end_idx > 0 and end_idx < len(raw_text) - 1:
            raw_text = raw_text[:end_idx + 1]
            logger.info(f"🧹 [Document Parser] Removed text after last brace")

        # Fix common JSON issues
        sanitized = raw_text

        # Remove trailing commas before closing brackets/braces
        sanitized = re.sub(r',\s*}', '}', sanitized)
        sanitized = re.sub(r',\s*]', ']', sanitized)

        # Fix newlines within strings (common LLM issue)
        sanitized = re.sub(r'"([^"]*)\n([^"]*)"', r'"\1 \2"', sanitized)

        logger.info(f"🧹 [Document Parser] Gentle sanitization complete. Length: {len(raw_text)} -> {len(sanitized)}")
        return sanitized

    def _extract_field_mappings_from_text(self, raw_text: str) -> List[Dict[str, Any]]:
        """Extract field mappings from raw text using regex patterns"""
        import re

        field_mappings = []

        # Look for field mapping patterns in the text
        # Pattern: "field": "value" with confidence
        field_patterns = [
            r'"field":\s*"([^"]+)",\s*"value":\s*"([^"]+)",\s*"confidence":\s*([\d.]+)',
            r'"field":\s*"([^"]+)",\s*"value":\s*"([^"]+)"',
            r'Field:\s*([^\n]+)\s*Value:\s*([^\n]+)\s*Confidence:\s*([\d.]+)',
        ]

        for pattern in field_patterns:
            matches = re.findall(pattern, raw_text, re.IGNORECASE)
            for match in matches:
                if len(match) == 3:
                    field, value, confidence = match
                    field_mappings.append({
                        "field": field.strip(),
                        "value": value.strip(),
                        "confidence": float(confidence),
                        "source": "regex_extraction",
                        "reasoning": "Extracted via pattern matching",
                        "priority": "medium"
                    })
                elif len(match) == 2:
                    field, value = match
                    field_mappings.append({
                        "field": field.strip(),
                        "value": value.strip(),
                        "confidence": 0.6,
                        "source": "regex_extraction",
                        "reasoning": "Extracted via pattern matching",
                        "priority": "medium"
                    })

        return field_mappings

    def _validate_and_fix_rfq_structure(self, rfq_data: Dict[str, Any]) -> None:
        """Validate and fix RFQ data structure"""
        # Ensure required keys exist
        if "field_mappings" not in rfq_data:
            rfq_data["field_mappings"] = []

        if "confidence" not in rfq_data:
            rfq_data["confidence"] = 0.8

        if "identified_sections" not in rfq_data:
            rfq_data["identified_sections"] = {}

        if "extracted_entities" not in rfq_data:
            rfq_data["extracted_entities"] = {
                "stakeholders": [],
                "industries": [],
                "research_types": [],
                "methodologies": []
            }

        # Validate field mappings structure
        valid_mappings = []
        for mapping in rfq_data.get("field_mappings", []):
            if isinstance(mapping, dict) and "field" in mapping and "value" in mapping:
                # Ensure required fields exist
                if "confidence" not in mapping:
                    mapping["confidence"] = 0.6
                if "source" not in mapping:
                    mapping["source"] = "llm_extraction"
                if "reasoning" not in mapping:
                    mapping["reasoning"] = "Extracted by LLM"
                if "priority" not in mapping:
                    mapping["priority"] = "medium"
                valid_mappings.append(mapping)

        rfq_data["field_mappings"] = valid_mappings
        logger.info(f"✅ [Document Parser] Validated RFQ structure with {len(valid_mappings)} field mappings")

    def _get_fallback_rfq_structure(self, error_message: str) -> Dict[str, Any]:
        """Return fallback RFQ structure when parsing fails"""
        return {
            "confidence": 0.0,
            "identified_sections": {},
            "extracted_entities": {
                "stakeholders": [],
                "industries": [],
                "research_types": [],
                "methodologies": []
            },
            "field_mappings": [],
            "extraction_error": f"RFQ extraction failed: {error_message}"
        }

    async def parse_document_for_rfq(self, docx_content: bytes, filename: str = None, session_id: str = None) -> Dict[str, Any]:
        """Parse DOCX document specifically for RFQ data extraction."""
        logger.info(f"🎯 [Document Parser] Starting RFQ-specific document parsing")

        # Import progress tracker
        from .progress_tracker import get_progress_tracker
        tracker = get_progress_tracker(session_id or "document_parse")

        try:
            # Send initial progress
            progress_data = tracker.get_progress_data("extracting_document")
            await self._send_progress(session_id, "extracting", progress_data["percent"], progress_data["message"])

            # Extract text from DOCX
            logger.info(f"📄 [Document Parser] Extracting text from DOCX")
            document_text = await self.extract_text_from_docx(docx_content, session_id)

            if not document_text.strip():
                logger.error(f"❌ [Document Parser] No text content found in document")
                await self._send_progress(session_id, "error", 0, "No text content found in document")
                raise DocumentParsingError("No text content found in document")

            logger.info(f"✅ [Document Parser] Text extraction successful, length: {len(document_text)} chars")

            # Send progress update for document processing
            progress_data = tracker.get_progress_data("processing_document")
            await self._send_progress(session_id, "prompting", progress_data["percent"], progress_data["message"])

            # Extract RFQ-specific data
            logger.info(f"🎯 [Document Parser] Extracting RFQ data from text")
            progress_data = tracker.get_progress_data("analyzing_document")
            await self._send_progress(session_id, "llm_processing", progress_data["percent"], progress_data["message"])

            rfq_data = await self.extract_rfq_data(document_text)
            logger.info(f"✅ [Document Parser] RFQ data extraction completed")

            # Send progress for parsing completion
            progress_data = tracker.get_progress_data("parsing_complete")
            await self._send_progress(session_id, "parsing", progress_data["percent"], progress_data["message"])

            # Structure the response for frontend consumption
            result = {
                "document_content": {
                    "raw_text": document_text,
                    "filename": filename,
                    "word_count": len(document_text.split()),
                    "extraction_timestamp": "2024-01-01T00:00:00Z"
                },
                "rfq_analysis": rfq_data,
                "processing_status": "completed",
                "errors": []
            }

            logger.info(f"🎉 [Document Parser] RFQ document parsing completed successfully")
            logger.info(f"📊 [Document Parser] Extraction confidence: {rfq_data.get('confidence', 0)}")
            logger.info(f"📊 [Document Parser] Field mappings found: {len(rfq_data.get('field_mappings', []))}")

            # Send completion progress
            completion_data = tracker.get_completion_data("parsing_complete")
            await self._send_progress(session_id, "completed", completion_data["percent"], completion_data["message"])

            return result

        except DocumentParsingError:
            await self._send_progress(session_id, "error", 0, "Document parsing failed")
            raise
        except Exception as e:
            logger.error(f"❌ [Document Parser] Unexpected error during RFQ document parsing: {str(e)}", exc_info=True)
            await self._send_progress(session_id, "error", 0, f"Unexpected error: {str(e)}")
            raise DocumentParsingError(f"Unexpected error: {str(e)}")

    async def parse_document(self, docx_content: bytes) -> Dict[str, Any]:
        """Main method to parse DOCX document and return validated JSON."""
        logger.info(f"🤖 [Document Parser] Starting main document parsing process")
        try:
            # Extract text from DOCX
            logger.info(f"📄 [Document Parser] Extracting text from DOCX")
            document_text = await self.extract_text_from_docx(docx_content)

            if not document_text.strip():
                logger.error(f"❌ [Document Parser] No text content found in document")
                raise DocumentParsingError("No text content found in document")

            logger.info(f"✅ [Document Parser] Text extraction successful, length: {len(document_text)} chars")

            # NEW: Extract comments from DOCX
            logger.info(f"💬 [Document Parser] Extracting comments from DOCX")
            comments = self.extract_comments_from_docx(docx_content)
            logger.info(f"✅ [Document Parser] Found {len(comments)} comments")

            # Convert to JSON using LLM with comment context
            logger.info(f"🤖 [Document Parser] Converting text to JSON using LLM with comment context")
            survey_json = await self.convert_to_json(document_text, comments)
            logger.info(f"✅ [Document Parser] LLM conversion completed")

            # Validate the JSON
            logger.info(f"🔍 [Document Parser] Validating survey JSON structure")
            validated_survey = self.validate_survey_json(survey_json)
            logger.info(f"✅ [Document Parser] JSON validation successful")

            # NEW: Create question annotations from comments
            # Check multiple possible locations for questions
            questions_to_process = None
            if comments:
                # Check direct questions array
                if "questions" in validated_survey and validated_survey["questions"]:
                    questions_to_process = validated_survey["questions"]
                    logger.info(f"💬 [Document Parser] Found questions directly in survey_json")
                    
                    # Post-process questions in direct array
                    logger.info(f"🔧 [Document Parser] Post-processing {len(questions_to_process)} questions from direct array...")
                    for i, question in enumerate(questions_to_process):
                        if isinstance(question, dict):
                            processed_question = self._post_process_question(question)
                            validated_survey["questions"][i] = processed_question
                
                # Check final_output questions
                elif "final_output" in validated_survey and "questions" in validated_survey["final_output"] and validated_survey["final_output"]["questions"]:
                    questions_to_process = validated_survey["final_output"]["questions"]
                    logger.info(f"💬 [Document Parser] Found questions in final_output")
                    
                    # Post-process questions in final_output
                    logger.info(f"🔧 [Document Parser] Post-processing {len(questions_to_process)} questions from final_output...")
                    for i, question in enumerate(questions_to_process):
                        if isinstance(question, dict):
                            processed_question = self._post_process_question(question)
                            validated_survey["final_output"]["questions"][i] = processed_question
                
                # Check sections for questions (NEW: Handle sections format)
                elif "sections" in validated_survey and validated_survey["sections"]:
                    all_questions = []
                    for section in validated_survey["sections"]:
                        if isinstance(section, dict) and "questions" in section and section["questions"]:
                            all_questions.extend(section["questions"])
                    
                    # Post-process all questions in sections
                    logger.info(f"🔧 [Document Parser] Post-processing {len(all_questions)} questions from sections...")
                    for question in all_questions:
                        if isinstance(question, dict):
                            processed_question = self._post_process_question(question)
                            # Update the question in the section
                            for section in validated_survey["sections"]:
                                if isinstance(section, dict) and "questions" in section:
                                    for i, q in enumerate(section["questions"]):
                                        if q == question:
                                            section["questions"][i] = processed_question
                                            break

                    # Post-process textBlocks in sections
                    logger.info(f"🔧 [Document Parser] Post-processing textBlocks in sections...")
                    for section in validated_survey.get("sections", []):
                        if isinstance(section, dict) and "textBlocks" in section:
                            processed_text_blocks = []
                            for text_block in section["textBlocks"]:
                                if isinstance(text_block, dict):
                                    processed_text_block = self._post_process_text_block(text_block)
                                    processed_text_blocks.append(processed_text_block)
                                else:
                                    processed_text_blocks.append(text_block)
                            section["textBlocks"] = processed_text_blocks

            # Check raw_output for questions (for golden pairs)
            if "raw_output" in validated_survey and "questions" in validated_survey["raw_output"] and validated_survey["raw_output"]["questions"]:
                questions_to_process = validated_survey["raw_output"]["questions"]
                logger.info(f"💬 [Document Parser] Found questions in raw_output")
            
            if not questions_to_process:
                logger.warning(f"⚠️ [Document Parser] No questions found in any expected location")
            
            if questions_to_process:
                survey_id = validated_survey.get("survey_id", str(uuid.uuid4()))
                logger.info(f"💬 [Document Parser] Creating question annotations from comments for {len(questions_to_process)} questions")
                annotations_created = self.create_question_annotations_from_comments(
                    survey_id, 
                    questions_to_process, 
                    comments
                )
                logger.info(f"✅ [Document Parser] Question annotations created successfully")
                
                # Sync annotations to RAG if any were created
                if annotations_created > 0:
                    logger.info(f"🔗 [Document Parser] Syncing {annotations_created} annotations to RAG")
                    try:
                        from src.services.annotation_rag_sync_service import AnnotationRAGSyncService
                        
                        sync_service = AnnotationRAGSyncService(self.db_session)
                        
                        # Get all annotations we just created
                        from ..database.models import QuestionAnnotation
                        recent_annotations = self.db_session.query(QuestionAnnotation).filter(
                            QuestionAnnotation.survey_id == survey_id,
                            QuestionAnnotation.annotator_id == "docx_parser"
                        ).all()
                        
                        sync_count = 0
                        for annotation in recent_annotations:
                            result = await sync_service.sync_question_annotation(annotation.id)
                            if result.get("success"):
                                sync_count += 1
                                logger.info(f"✅ Synced annotation to RAG: {result.get('action')}")
                        
                        logger.info(f"🎉 [Document Parser] Synced {sync_count}/{len(recent_annotations)} annotations to RAG")
                        
                    except Exception as e:
                        logger.warning(f"⚠️ [Document Parser] Failed to sync annotations to RAG (non-critical): {str(e)}")

            # NEW: Add comment metadata to survey JSON
            if comments:
                # Extract unique comment categories
                comment_categories = list(set(comment["text"] for comment in comments))
                
                validated_survey["comment_metadata"] = {
                    "total_comments": len(comments),
                    "comment_categories": comment_categories,
                    "extraction_timestamp": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime())
                }
                logger.info(f"📊 [Document Parser] Added comment metadata: {len(comments)} comments, {len(comment_categories)} unique categories")

            # Log final survey structure
            logger.info(f"📊 [Document Parser] Final survey structure: {list(validated_survey.keys()) if isinstance(validated_survey, dict) else 'Not a dict'}")
            if isinstance(validated_survey, dict):
                logger.info(f"📊 [Document Parser] Survey title: {validated_survey.get('title', 'No title')}")
                logger.info(f"📊 [Document Parser] Questions count: {len(validated_survey.get('questions', []))}")
                logger.info(f"📊 [Document Parser] Confidence score: {validated_survey.get('confidence_score', 'No score')}")

            logger.info(f"🎉 [Document Parser] Document parsing completed successfully")
            return validated_survey

        except DocumentParsingError:
            raise
        except Exception as e:
            logger.error(f"❌ [Document Parser] Unexpected error during document parsing: {str(e)}", exc_info=True)
            raise DocumentParsingError(f"Unexpected error: {str(e)}")

    def _extract_comment_ranges_from_document(self, docx_content: bytes) -> Dict[str, str]:
        """
        Extract the exact text that each comment is anchored to using commentRangeStart/End markers.
        Returns a mapping of comment_id -> anchored_text.
        """
        import zipfile
        import xml.etree.ElementTree as ET
        
        comment_ranges = {}
        
        try:
            with zipfile.ZipFile(BytesIO(docx_content), 'r') as docx_zip:
                if 'word/document.xml' not in docx_zip.namelist():
                    logger.warning(f"⚠️ [Comment Ranges] No document.xml found")
                    return comment_ranges
                
                document_xml = docx_zip.read('word/document.xml')
                doc_root = ET.fromstring(document_xml)
                
                # Namespace for Word XML
                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                
                # Find all paragraphs and process them
                for para in doc_root.findall('.//w:p', ns):
                    # Track if we're inside a comment range
                    current_comment_id = None
                    anchored_text = []
                    
                    # Process all elements in the paragraph
                    for elem in para.iter():
                        # Check for comment range start
                        if elem.tag == f"{{{ns['w']}}}commentRangeStart":
                            current_comment_id = elem.get(f"{{{ns['w']}}}id")
                            anchored_text = []
                        
                        # Collect text if we're inside a comment range
                        elif current_comment_id and elem.tag == f"{{{ns['w']}}}t":
                            if elem.text:
                                anchored_text.append(elem.text)
                        
                        # Check for comment range end
                        elif elem.tag == f"{{{ns['w']}}}commentRangeEnd":
                            end_comment_id = elem.get(f"{{{ns['w']}}}id")
                            if end_comment_id == current_comment_id and anchored_text:
                                # Store the anchored text for this comment
                                comment_ranges[current_comment_id] = ''.join(anchored_text).strip()
                                logger.debug(f"🔗 [Comment Ranges] Comment {current_comment_id} anchored to: '{comment_ranges[current_comment_id][:50]}...'")
                            current_comment_id = None
                            anchored_text = []
                
                logger.info(f"✅ [Comment Ranges] Extracted {len(comment_ranges)} comment anchor ranges")
                return comment_ranges
                
        except Exception as e:
            logger.error(f"❌ [Comment Ranges] Error extracting comment ranges: {str(e)}")
            return comment_ranges

    def extract_comments_from_docx(self, docx_content: bytes) -> List[Dict[str, Any]]:
        """
        Extract comments from a DOCX file with exact anchored text from commentRange markers.
        Returns a list of comment dictionaries with author, date, text, and anchored_text.
        """
        import zipfile
        import xml.etree.ElementTree as ET
        
        try:
            comments = []
            
            with zipfile.ZipFile(BytesIO(docx_content), 'r') as docx_zip:
                # Check if comments.xml exists
                if 'word/comments.xml' not in docx_zip.namelist():
                    logger.info(f"📄 [Comment Extraction] No comments.xml found in DOCX")
                    return comments
                
                # Read comments.xml
                comments_xml = docx_zip.read('word/comments.xml')
                root = ET.fromstring(comments_xml)
                
                # Find all comment elements
                for comment_elem in root.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}comment'):
                    comment_id = comment_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                    author = comment_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author')
                    date = comment_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date')
                    
                    # Extract comment text
                    comment_text = ""
                    for p in comment_elem.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
                        for t in p.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                            if t.text:
                                comment_text += t.text
                        comment_text += "\n"
                    
                    comments.append({
                        'id': comment_id,
                        'author': author,
                        'date': date,
                        'text': comment_text.strip()
                    })
            
            # Extract exact anchored text ranges from document.xml
            logger.info(f"🔗 [Comment Extraction] Extracting comment anchor ranges from document.xml")
            comment_ranges = self._extract_comment_ranges_from_document(docx_content)
            
            # Merge anchored text with comment data
            for comment in comments:
                comment_id = comment['id']
                if comment_id in comment_ranges:
                    comment['anchored_text'] = comment_ranges[comment_id]
                    logger.info(f"✅ [Comment Extraction] Comment {comment_id} anchored to: '{comment['anchored_text'][:50]}...'")
                else:
                    comment['anchored_text'] = None
                    logger.warning(f"⚠️ [Comment Extraction] No anchor text found for comment {comment_id}")
            
            logger.info(f"✅ [Comment Extraction] Found {len(comments)} comments, {len(comment_ranges)} with anchored text")
            return comments
            
        except Exception as e:
            logger.error(f"❌ [Comment Extraction] Error extracting comments: {str(e)}")
            return []

    def _add_positional_context_to_comments(self, docx_content: bytes, comments: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        Add positional context to comments by parsing the main document structure.
        This helps match comments to the actual content they're commenting on.
        """
        import zipfile
        import xml.etree.ElementTree as ET
        
        try:
            comments_with_context = []
            
            with zipfile.ZipFile(BytesIO(docx_content), 'r') as docx_zip:
                # Read document.xml to get the main content structure
                if 'word/document.xml' not in docx_zip.namelist():
                    logger.warning(f"⚠️ [Comment Context] No document.xml found")
                    return comments
                
                document_xml = docx_zip.read('word/document.xml')
                doc_root = ET.fromstring(document_xml)
                
                # Extract all paragraphs with their content
                paragraphs = []
                for p_elem in doc_root.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
                    paragraph_text = ""
                    for t in p_elem.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                        if t.text:
                            paragraph_text += t.text
                    
                    if paragraph_text.strip():
                        paragraphs.append(paragraph_text.strip())
                
                # For each comment, try to find the closest paragraph content
                for comment in comments:
                    comment_text = comment.get('text', '').lower()
                    comment_with_context = comment.copy()
                    
                    # Find the best matching paragraph based on content similarity
                    best_match_score = 0
                    best_match_context = ""
                    best_match_position = 0
                    
                    for i, paragraph in enumerate(paragraphs):
                        paragraph_lower = paragraph.lower()
                        
                        # Calculate similarity score based on common words
                        comment_words = set(comment_text.split())
                        paragraph_words = set(paragraph_lower.split())
                        
                        if comment_words and paragraph_words:
                            common_words = comment_words.intersection(paragraph_words)
                            similarity_score = len(common_words) / len(comment_words)
                            
                            if similarity_score > best_match_score:
                                best_match_score = similarity_score
                                best_match_context = paragraph
                                best_match_position = i
                    
                    # Add context information to the comment
                    comment_with_context.update({
                        'context': best_match_context,
                        'context_position': best_match_position,
                        'context_similarity': best_match_score,
                        'surrounding_paragraphs': paragraphs[max(0, best_match_position-1):best_match_position+2] if paragraphs else []
                    })
                    
                    comments_with_context.append(comment_with_context)
                    
                    logger.debug(f"🔍 [Comment Context] Comment '{comment_text[:30]}...' matched to paragraph {best_match_position} with score {best_match_score:.2f}")
            
            return comments_with_context
            
        except Exception as e:
            logger.error(f"❌ [Comment Context] Error adding positional context: {str(e)}")
            return comments

    def find_best_comment_match(self, question: Dict[str, Any], comments: List[Dict[str, Any]]) -> Optional[Dict[str, Any]]:
        """
        Find the best matching comment for a question using contextual analysis.
        Uses positional context, content similarity, and semantic matching.
        """
        if not comments:
            return None
        
        question_text = question.get('text', '').lower()
        question_id = question.get('id', '')
        
        best_match = None
        best_score = 0
        
        for comment in comments:
            comment_text = comment.get('text', '').lower()
            context = comment.get('context', '').lower()
            context_similarity = comment.get('context_similarity', 0)
            
            # Calculate multiple matching scores
            scores = []
            
            # 1. Direct text similarity between question and comment
            if question_text and comment_text:
                question_words = set(question_text.split())
                comment_words = set(comment_text.split())
                if question_words and comment_words:
                    direct_similarity = len(question_words.intersection(comment_words)) / len(question_words)
                    scores.append(('direct', direct_similarity))
            
            # 2. Context similarity (how well comment matches its surrounding content)
            scores.append(('context', context_similarity))
            
            # 3. Semantic keyword matching
            semantic_score = self._calculate_semantic_similarity(question_text, comment_text)
            scores.append(('semantic', semantic_score))
            
            # 4. Question ID pattern matching (e.g., SQ01, AQ01, BQ01)
            pattern_score = self._calculate_pattern_similarity(question_id, comment_text)
            scores.append(('pattern', pattern_score))
            
            # Calculate weighted overall score
            weights = {'direct': 0.4, 'context': 0.3, 'semantic': 0.2, 'pattern': 0.1}
            overall_score = sum(score * weights.get(score_type, 0) for score_type, score in scores)
            
            logger.debug(f"🔍 [Comment Matching] Question '{question_id}' vs Comment '{comment_text[:30]}...': {dict(scores)} = {overall_score:.3f}")
            
            if overall_score > best_score:
                best_score = overall_score
                best_match = comment
        
        # Only return a match if the score is above a threshold
        if best_score > 0.1:  # Minimum threshold for matching
            logger.info(f"✅ [Comment Matching] Best match for question '{question_id}': '{best_match['text'][:50]}...' (score: {best_score:.3f})")
            return best_match
        else:
            logger.debug(f"⚠️ [Comment Matching] No suitable match found for question '{question_id}' (best score: {best_score:.3f})")
            return None

    def _calculate_semantic_similarity(self, question_text: str, comment_text: str) -> float:
        """Calculate semantic similarity based on research methodology keywords."""
        methodology_keywords = {
            'screening': ['screen', 'qualify', 'eligibility', 'criteria'],
            'demographic': ['age', 'gender', 'income', 'education', 'demographic'],
            'usage': ['use', 'usage', 'behavior', 'frequency', 'experience'],
            'product': ['product', 'brand', 'awareness', 'satisfaction', 'preference'],
            'concept': ['concept', 'impression', 'likelihood', 'purchase', 'intent'],
            'pricing': ['price', 'cost', 'value', 'willingness', 'pay'],
            'conjoint': ['choice', 'scenario', 'attribute', 'preference'],
            'maxdiff': ['important', 'least', 'most', 'feature', 'benefit']
        }
        
        question_lower = question_text.lower()
        comment_lower = comment_text.lower()
        
        max_category_score = 0
        for category, keywords in methodology_keywords.items():
            question_matches = sum(1 for keyword in keywords if keyword in question_lower)
            comment_matches = sum(1 for keyword in keywords if keyword in comment_lower)
            
            if question_matches > 0 and comment_matches > 0:
                category_score = min(question_matches, comment_matches) / max(question_matches, comment_matches)
                max_category_score = max(max_category_score, category_score)
        
        return max_category_score

    def _calculate_pattern_similarity(self, question_id: str, comment_text: str) -> float:
        """Calculate similarity based on question ID patterns and comment content."""
        if not question_id or not comment_text:
            return 0
        
        comment_lower = comment_text.lower()
        
        # Extract question prefix (SQ, AQ, BQ, etc.)
        question_prefix = ''.join([c for c in question_id if c.isalpha()])
        
        # Map question prefixes to likely comment categories
        prefix_mapping = {
            'SQ': ['screen', 'qualify', 'eligibility', 'intro'],
            'AQ': ['awareness', 'usage', 'behavior', 'current'],
            'BQ': ['brand', 'product', 'concept', 'impression'],
            'CQ': ['choice', 'conjoint', 'scenario'],
            'DQ': ['demographic', 'profile', 'background']
        }
        
        expected_keywords = prefix_mapping.get(question_prefix, [])
        if expected_keywords:
            matches = sum(1 for keyword in expected_keywords if keyword in comment_lower)
            return matches / len(expected_keywords)
        
        return 0

    def create_question_annotations_from_comments(
        self, 
        survey_id: str, 
        questions: List[Dict], 
        comments: List[Dict]
    ) -> None:
        """Create question annotations from LLM-embedded annotation data."""
        
        if not self.db_session:
            logger.warning(f"⚠️ [Comment Annotation] No database session available, skipping annotation creation")
            return
        
        try:
            from ..database.models import QuestionAnnotation
            from src.services.label_normalizer import LabelNormalizer
            
            logger.info(f"💬 [Comment Annotation] Processing annotations for {len(questions)} questions")
            
            # Initialize label normalizer
            label_normalizer = LabelNormalizer()
            
            annotations_created = 0
            
            for i, question in enumerate(questions):
                question_id = question.get("id", f"q_{i+1}")
                
                # Check if this question has an annotation embedded by the LLM
                if "annotation" in question and isinstance(question["annotation"], dict):
                    annotation_data = question["annotation"]
                    
                    # Extract annotation details
                    comment_text = annotation_data.get("comment", "")
                    anchored_text = annotation_data.get("anchored_text", "")
                    author = annotation_data.get("author", "docx_parser")
                    date = annotation_data.get("date", "")
                    
                    if comment_text:  # Only create annotation if there's actual comment text
                        # Normalize the comment text as a label
                        normalized_label = label_normalizer.normalize(comment_text)
                        
                        logger.info(f"🏷️ [Comment Annotation] Normalizing comment '{comment_text}' -> '{normalized_label}'")
                        
                        # Create advanced_labels with full context
                        advanced_labels = {
                            "comment_text": comment_text,
                            "anchored_text": anchored_text,
                            "comment_author": author,
                            "comment_date": date,
                            "matching_method": "llm_with_anchored_text",
                            "matching_confidence": 1.0,  # LLM did the matching, so high confidence
                            "original_comment": comment_text,
                            "normalized_label": normalized_label
                        }
                        
                        # Create annotation with normalized labels
                        annotation = QuestionAnnotation(
                            question_id=question_id,
                            survey_id=survey_id,
                            
                            # Store both original and normalized
                            comment=comment_text,
                            labels=[normalized_label],  # Use normalized label
                            
                            # Store full context in advanced_labels
                            advanced_labels=advanced_labels,
                            
                            # Use comment author as annotator_id
                            annotator_id=author,
                            
                            # Default values for required fields
                            required=True,
                            quality=3,
                            relevant=3,
                            methodological_rigor=3,
                            content_validity=3,
                            respondent_experience=3,
                            analytical_value=3,
                            business_impact=3,
                            
                            # Mark as AI-generated from DOCX parsing
                            ai_generated=True,
                            ai_confidence=1.0,
                            generation_timestamp=time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime())
                        )
                        
                        self.db_session.add(annotation)
                        annotations_created += 1
                        logger.info(f"✅ [Comment Annotation] Created annotation for question {question_id}: '{comment_text[:50]}...'")
                    else:
                        logger.debug(f"⚠️ [Comment Annotation] Question {question_id} has annotation field but no comment text")
                else:
                    logger.debug(f"ℹ️ [Comment Annotation] Question {question_id} has no annotation field")
            
            self.db_session.commit()
            logger.info(f"🎉 [Comment Annotation] Successfully created {annotations_created} question annotations from LLM-embedded data")
            
            return annotations_created
            
        except Exception as e:
            logger.error(f"❌ [Comment Annotation] Error creating annotations: {str(e)}")
            if self.db_session:
                self.db_session.rollback()
            raise

# Global instance
document_parser = DocumentParser()