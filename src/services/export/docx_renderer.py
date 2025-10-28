"""
DOCX export renderer for surveys.
Implements the abstract base renderer with Word document generation.
"""

import re
from io import BytesIO
from typing import Dict, List, Any

try:
    from docx import Document
    from docx.document import Document as DocumentType
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    # DocX not available - used for testing environments
    Document = None
    DocumentType = None
    Inches = None
    Pt = None
    WD_ALIGN_PARAGRAPH = None
    WD_TABLE_ALIGNMENT = None
    OxmlElement = None
    qn = None
    DOCX_AVAILABLE = False

from .base import SurveyExportRenderer, QuestionType, export_registry


class DocxSurveyRenderer(SurveyExportRenderer):
    """
    DOCX renderer for surveys using python-docx.
    Implements all question types with appropriate Word formatting.
    """

    def __init__(self) -> None:
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DocxSurveyRenderer but not available")
        self.document: DocumentType | None = None
        super().__init__()

    def _register_question_types(self) -> None:
        """Register all implemented question types."""
        self._registered_types = {t.value for t in QuestionType}

    def _get_document(self) -> DocumentType:
        """Get the document instance with type assertion."""
        assert self.document is not None, "Document not initialized"
        return self.document

    def _initialize_document(self, survey_data: Dict[str, Any]) -> None:
        """Initialize the Word document."""
        self.document = Document()

        # Set document margins
        sections = self._get_document().sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    def _render_survey_header(self, title: str, description: str = "") -> None:
        """Render survey title and description."""
        doc = self._get_document()
        # Title
        title_paragraph = doc.add_heading(title, level=1)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Description
        if description:
            desc_paragraph = doc.add_paragraph(description)
            desc_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add some space
        doc.add_paragraph("")

    def render_survey(self, survey_data: Dict[str, Any]) -> bytes:
        """
        Main entry point for rendering a complete survey with sections support.

        Args:
            survey_data: Dictionary containing survey metadata and sections/questions

        Returns:
            Rendered survey as bytes
        """
        self._initialize_document(survey_data)

        # Render survey header/metadata
        if "title" in survey_data:
            self._render_survey_header(survey_data["title"], survey_data.get("description", ""))

        # Check if survey uses sections format
        if "sections" in survey_data:
            # Render sections with text blocks and questions
            sections = survey_data.get("sections", [])
            global_question_number = 1
            for section in sections:
                self._render_section(section, global_question_number)
                # Update global question number after processing section
                questions = section.get("questions", [])
                for question in questions:
                    if question.get("type") != "instruction":
                        global_question_number += 1
        else:
            # Legacy format: render questions directly
            questions = survey_data.get("questions", [])
            for question in questions:
                self._render_question(question)

        return self._finalize_document()

    def _finalize_document(self) -> bytes:
        """Finalize and return the document as bytes."""
        buffer = BytesIO()
        self._get_document().save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _add_question_header(self, question: Dict[str, Any], question_number: int | None = None) -> None:
        """Add question header with numbering and text."""
        doc = self._get_document()
        question_text = question.get("text", "")

        if question_number:
            header_text = f"Q{question_number}: {question_text}"
        else:
            header_text = question_text

        paragraph = doc.add_paragraph()
        run = paragraph.add_run(header_text)
        run.bold = True
        run.font.size = Pt(11)

    def _add_required_indicator(self, question: Dict[str, Any]) -> None:
        """Add required field indicator if applicable."""
        if question.get("required", False):
            doc = self._get_document()
            paragraph = doc.add_paragraph()
            run = paragraph.add_run("* Required")
            run.italic = True
            run.font.size = Pt(9)

    def _render_section(self, section: Dict[str, Any], start_question_number: int = 1) -> None:
        """Render a survey section with text blocks and questions."""
        doc = self._get_document()

        # Add section header
        section_title = section.get("title", f"Section {section.get('id', 'Unknown')}")
        section_heading = doc.add_heading(section_title, level=2)
        section_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add section description if available
        if section.get("description"):
            desc_paragraph = doc.add_paragraph(section["description"])
            desc_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Render intro text if available
        if section.get("introText"):
            self._render_text_block(section["introText"], "Introduction")
        
        # Render text blocks if available (sorted by order)
        text_blocks = section.get("textBlocks", [])
        for text_block in sorted(text_blocks, key=lambda x: x.get("order", 0)):
            self._render_text_block(text_block, "Information")
        
        # Render questions if available
        questions = section.get("questions", [])
        question_number = start_question_number
        for question in questions:
            if question.get("type") == "instruction":
                # Instruction questions don't get numbered
                self._render_instruction(question)
            else:
                # Regular questions get numbered
                self._render_question_with_number(question, question_number)
                question_number += 1
        
        # Render closing text if available
        if section.get("closingText"):
            self._render_text_block(section["closingText"], "Closing")
        
        # Add space between sections
        doc.add_paragraph("")

    def _render_text_block(self, text_block: Dict[str, Any], block_type: str = "Information") -> None:
        """Render a text block (introText, textBlocks, closingText)."""
        doc = self._get_document()
        
        # Get text block properties
        content = text_block.get("content", "")
        text_type = text_block.get("type", "information")
        label = text_block.get("label", "")
        mandatory = text_block.get("mandatory", False)
        
        # Create a bordered paragraph for the text block
        paragraph = doc.add_paragraph()
        
        # Use label as main heading if available, otherwise use block_type
        main_heading = label if label else block_type
        if mandatory:
            main_heading += " [REQUIRED]"
        
        # Main heading (prominent)
        header_run = paragraph.add_run(main_heading)
        header_run.bold = True
        header_run.font.size = Pt(12)
        
        # Add subscript type label if different from main heading
        if label and block_type != "Information":
            subscript_paragraph = doc.add_paragraph()
            subscript_run = subscript_paragraph.add_run(block_type)
            subscript_run.italic = True
            subscript_run.font.size = Pt(8)
        
        # Add content
        content_paragraph = doc.add_paragraph()
        content_run = content_paragraph.add_run(content)
        content_run.font.size = Pt(9)
        
        # Add some space after text block
        doc.add_paragraph("")

    # Question type implementations

    def _render_instruction(self, question: Dict[str, Any]) -> None:
        """Render instruction question with first words before colon as heading."""
        doc = self._get_document()
        
        # Extract first words before colon as heading
        text = question.get("text", "")
        colon_index = text.find(':')
        
        if colon_index > 0:
            heading = text[:colon_index].strip()
            content = text[colon_index + 1:].strip()
        else:
            heading = "Instructions"
            content = text
        
        # Add heading (first words before colon)
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(heading)
        run.bold = True
        run.font.size = Pt(12)

        # Add content
        content_paragraph = doc.add_paragraph()
        content_run = content_paragraph.add_run(content)
        content_run.font.size = Pt(10)


        doc.add_paragraph("")

    def _render_question_with_number(self, question: Dict[str, Any], question_number: int) -> None:
        """Render a question with a specific question number."""
        # Store the question number temporarily
        original_method = self._add_question_header
        
        def _add_question_header_with_number(question: Dict[str, Any], question_number_param: int | None = None) -> None:
            # Use the provided question number
            original_method(question, question_number)
        
        # Temporarily replace the method
        self._add_question_header = _add_question_header_with_number
        
        try:
            # Render the question
            self._render_question(question)
        finally:
            # Restore the original method
            self._add_question_header = original_method

    def _render_multiple_choice(self, question: Dict[str, Any]) -> None:
        """Render multiple choice question with checkboxes."""
        self._add_question_header(question)
        self._add_required_indicator(question)

        doc = self._get_document()
        paragraph = doc.add_paragraph()
        paragraph.add_run("Select one or more options:").italic = True

        options = question.get("options", [])
        for option in options:
            option_paragraph = doc.add_paragraph()
            option_paragraph.add_run("☐ " + option)

        doc.add_paragraph("")

    def _render_single_choice(self, question: Dict[str, Any]) -> None:
        """Render single choice question with radio buttons."""
        self._add_question_header(question)
        self._add_required_indicator(question)

        doc = self._get_document()
        paragraph = doc.add_paragraph()
        paragraph.add_run("Select one option:").italic = True

        options = question.get("options", [])
        for option in options:
            option_paragraph = doc.add_paragraph()
            option_paragraph.add_run("○ " + option)

        doc.add_paragraph("")

    def _render_yes_no(self, question: Dict[str, Any]) -> None:
        """Render yes/no question with radio buttons."""
        self._add_question_header(question)
        self._add_required_indicator(question)

        doc = self._get_document()
        paragraph = doc.add_paragraph()
        paragraph.add_run("Select one option:").italic = True

        # Use provided options or default to Yes/No
        options = question.get("options", ["Yes", "No"])
        for option in options:
            option_paragraph = doc.add_paragraph()
            option_paragraph.add_run("○ " + option)

        doc.add_paragraph("")

    def _render_dropdown(self, question: Dict[str, Any]) -> None:
        """Render dropdown question as single choice with note."""
        self._add_question_header(question)
        self._add_required_indicator(question)

        doc = self._get_document()
        paragraph = doc.add_paragraph()
        paragraph.add_run("Select one option from dropdown:").italic = True

        options = question.get("options", [])
        for option in options:
            option_paragraph = doc.add_paragraph()
            option_paragraph.add_run("• " + option)

        # Add note about dropdown format
        note_paragraph = doc.add_paragraph()
        note_paragraph.add_run("(Note: This will be rendered as a dropdown in the actual survey)").italic = True
        note_paragraph.runs[0].font.size = Pt(9)

        doc.add_paragraph("")

    def _render_scale(self, question: Dict[str, Any]) -> None:
        """Render scale question with numbered options."""
        self._add_question_header(question)
        self._add_required_indicator(question)

        scale_labels = question.get("scale_labels", {})
        options = question.get("options", [])

        if scale_labels:
            # Show scale with labels
            doc = self._get_document()
            paragraph = doc.add_paragraph()
            paragraph.add_run("Rate on the following scale:").italic = True

            for i, option in enumerate(options, 1):
                label = scale_labels.get(str(i), "")
                scale_paragraph = doc.add_paragraph()
                scale_paragraph.add_run(f"○ {i} - {option}" + (f" ({label})" if label else ""))
        else:
            # Simple numeric scale
            doc = self._get_document()
            paragraph = doc.add_paragraph()
            paragraph.add_run("Rate from 1 to " + str(len(options)) + ":").italic = True

            scale_paragraph = doc.add_paragraph()
            scale_text = " | ".join([f"○ {i}" for i in range(1, len(options) + 1)])
            scale_paragraph.add_run(scale_text)

        doc.add_paragraph("")

    def _render_ranking(self, question: Dict[str, Any]) -> None:
        """Render ranking question with numbered lines."""
        self._add_question_header(question)
        self._add_required_indicator(question)

        doc = self._get_document()
        paragraph = doc.add_paragraph()
        paragraph.add_run("Rank the following items in order of preference (1 = most preferred):").italic = True

        options = question.get("options", [])
        for option in options:
            rank_paragraph = doc.add_paragraph()
            rank_paragraph.add_run(f"___ {option}")

        doc.add_paragraph("")

    def _render_text(self, question: Dict[str, Any]) -> None:
        """Render text input question."""
        self._add_question_header(question)
        self._add_required_indicator(question)

        # Add response area
        doc = self._get_document()
        for _ in range(3):
            doc.add_paragraph("_" * 80)

        doc.add_paragraph("")

    def _render_open_text(self, question: Dict[str, Any]) -> None:
        """Render open text question (alias for text)."""
        self._render_text(question)

    def _render_numeric(self, question: Dict[str, Any]) -> None:
        """Render numeric input question."""
        self._add_question_header(question)
        self._add_required_indicator(question)

        doc = self._get_document()
        paragraph = doc.add_paragraph()
        paragraph.add_run("Enter a numeric value: _______________")

        doc.add_paragraph("")

    def _render_date(self, question: Dict[str, Any]) -> None:
        """Render date input question."""
        self._add_question_header(question)
        self._add_required_indicator(question)

        doc = self._get_document()
        paragraph = doc.add_paragraph()
        paragraph.add_run("Enter date (MM/DD/YYYY): _______________")

        doc.add_paragraph("")

    def _render_boolean(self, question: Dict[str, Any]) -> None:
        """Render yes/no question."""
        self._add_question_header(question)
        self._add_required_indicator(question)

        doc = self._get_document()
        paragraph = doc.add_paragraph()
        paragraph.add_run("○ Yes    ○ No")

        doc.add_paragraph("")

    def _render_multiple_select(self, question: Dict[str, Any]) -> None:
        """Render multiple select question (alias for multiple choice)."""
        self._render_multiple_choice(question)

    def _render_matrix(self, question: Dict[str, Any]) -> None:
        """Render basic matrix question."""
        self._add_question_header(question)
        self._add_required_indicator(question)

        # This is a simplified matrix - for complex matrices use matrix_likert
        doc = self._get_document()
        options = question.get("options", [])
        if len(options) < 2:
            doc.add_paragraph("Matrix format: Please specify rows and columns in question options.")
            return

        # Create simple table
        table = doc.add_table(rows=len(options), cols=2)
        table.style = 'Table Grid'

        for i, option in enumerate(options):
            table.cell(i, 0).text = option
            table.cell(i, 1).text = "☐"

        doc.add_paragraph("")

    def _extract_attributes_from_text(self, text: str) -> List[str]:
        """Extract attributes from question text (used by specialized components)."""
        # Look for comma-separated attributes after question mark, colon, or period
        search_text = text

        # Check for question mark first
        question_mark_index = text.find('?')
        if question_mark_index != -1:
            search_text = text[question_mark_index + 1:].strip()
        else:
            # Check for colon
            colon_index = text.find(':')
            if colon_index != -1:
                search_text = text[colon_index + 1:].strip()
            else:
                # Check for period (for statements like "Rate the following items. Item1, Item2, Item3")
                period_index = text.find('.')
                if period_index != -1:
                    search_text = text[period_index + 1:].strip()
                else:
                    # Fallback: if no punctuation found, try to find patterns like "at the following" or "priced at"
                    # This handles cases like "priced at the following per box of 6 monthly lenses. $30, $35, $40"
                    following_patterns = [
                        "at the following",
                        "priced at",
                        "the following",
                        "as follows"
                    ]
                    for pattern in following_patterns:
                        pattern_index = text.lower().find(pattern)
                        if pattern_index != -1:
                            # Look for the next period or end of string
                            next_period = text.find('.', pattern_index)
                            if next_period != -1:
                                search_text = text[next_period + 1:].strip()
                                break
                            else:
                                # If no period after pattern, take everything after the pattern
                                search_text = text[pattern_index + len(pattern):].strip()
                                break

        # Remove trailing period if present
        clean_text = search_text.rstrip('.')

        # Split by comma and clean up
        attributes = [attr.strip() for attr in clean_text.split(',') if attr.strip()]
        
        # Additional validation: if we found attributes, make sure they look reasonable
        if attributes and len(attributes) > 1:
            # Check if attributes look like they could be valid (not just single words)
            # This helps filter out cases where we might have parsed incorrectly
            return attributes
        
        # If no attributes found with the above methods, try a more aggressive approach
        # Look for any comma-separated list in the text
        if not attributes:
            import re
            # Look for patterns like "$30, $35, $40" or "Item1, Item2, Item3"
            comma_patterns = [
                r'\$\d+(?:\.\d{2})?(?:\s*,\s*\$?\d+(?:\.\d{2})?)+',  # Price patterns like $30, $35, $40
                r'[A-Za-z][A-Za-z0-9\s]*(?:\s*,\s*[A-Za-z][A-Za-z0-9\s]*)+',  # Word patterns
                r'[\w\s]+(?:\s*,\s*[\w\s]+)+'  # General word patterns
            ]
            
            for pattern in comma_patterns:
                matches = re.findall(pattern, text)
                if matches:
                    # Take the longest match (most likely to be the attribute list)
                    longest_match = max(matches, key=len)
                    attributes = [attr.strip() for attr in longest_match.split(',') if attr.strip()]
                    if len(attributes) > 1:
                        break
            
            # If still no attributes, try the simple approach: find the last comma-separated list
            if not attributes:
                last_comma = text.rfind(',')
                if last_comma != -1:
                    # Look backwards to find the start of the list
                    start = text.rfind('.', 0, last_comma)
                    if start == -1:
                        start = text.rfind(':', 0, last_comma)
                    if start == -1:
                        start = text.rfind('?', 0, last_comma)
                    
                    if start != -1:
                        list_text = text[start+1:].strip()
                        attributes = [attr.strip() for attr in list_text.split(',') if attr.strip()]
                        # Clean up any trailing periods
                        if attributes and attributes[-1].endswith('.'):
                            attributes[-1] = attributes[-1].rstrip('.')
        
        return attributes

    def _render_matrix_likert(self, question: Dict[str, Any]) -> None:
        """Render matrix likert question with proper table formatting."""
        self._add_question_header(question)
        self._add_required_indicator(question)

        # Extract attributes from question text
        attributes = self._extract_attributes_from_text(question.get("text", ""))
        options = question.get("options", ["Strongly Disagree", "Disagree", "Neutral", "Agree", "Strongly Agree"])

        doc = self._get_document()
        if not attributes:
            doc.add_paragraph("Matrix Likert: Unable to parse attributes from question text.")
            return

        # Create table with headers
        table = doc.add_table(rows=len(attributes) + 1, cols=len(options) + 1)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        table.cell(0, 0).text = "Items"
        for j, option in enumerate(options):
            table.cell(0, j + 1).text = option

        # Data rows
        for i, attribute in enumerate(attributes):
            table.cell(i + 1, 0).text = attribute
            for j in range(len(options)):
                table.cell(i + 1, j + 1).text = "○"

        doc.add_paragraph("")

    def _render_constant_sum(self, question: Dict[str, Any]) -> None:
        """Render constant sum allocation question."""
        self._add_question_header(question)
        self._add_required_indicator(question)

        # Parse total points from question text
        text = question.get("text", "")
        points_match = re.search(r'(\d+)\s+points?', text, re.IGNORECASE)
        total_points = int(points_match.group(1)) if points_match else 100

        # Extract attributes
        attributes = self._extract_attributes_from_text(text)

        doc = self._get_document()
        if not attributes:
            doc.add_paragraph("Constant Sum: Unable to parse attributes from question text.")
            return

        paragraph = doc.add_paragraph()
        paragraph.add_run(f"Allocate {total_points} points across the following items:").italic = True

        # Create allocation table
        table = doc.add_table(rows=len(attributes) + 1, cols=2)
        table.style = 'Table Grid'

        table.cell(0, 0).text = "Item"
        table.cell(0, 1).text = "Points"

        for i, attribute in enumerate(attributes):
            table.cell(i + 1, 0).text = attribute
            table.cell(i + 1, 1).text = "____"

        # Add total row
        total_paragraph = doc.add_paragraph()
        total_paragraph.add_run(f"Total: _____ / {total_points} points").bold = True

        doc.add_paragraph("")

    def _render_numeric_grid(self, question: Dict[str, Any]) -> None:
        """Render numeric grid question."""
        self._add_question_header(question)
        self._add_required_indicator(question)

        # Extract attributes
        attributes = self._extract_attributes_from_text(question.get("text", ""))
        options = question.get("options", [])

        doc = self._get_document()
        if not attributes:
            doc.add_paragraph("Numeric Grid: Unable to parse attributes from question text.")
            return

        paragraph = doc.add_paragraph()
        paragraph.add_run("Provide numeric values for each item:").italic = True

        # Create grid table
        table = doc.add_table(rows=len(attributes) + 1, cols=len(options) + 1)
        table.style = 'Table Grid'

        # Header row
        table.cell(0, 0).text = "Items"
        for j, option in enumerate(options):
            table.cell(0, j + 1).text = option

        # Data rows
        for i, attribute in enumerate(attributes):
            table.cell(i + 1, 0).text = attribute
            for j in range(len(options)):
                table.cell(i + 1, j + 1).text = "____"

        doc.add_paragraph("")

    def _render_numeric_open(self, question: Dict[str, Any]) -> None:
        """Render numeric open question (currency detected only when context suggests it)."""
        self._add_question_header(question)
        self._add_required_indicator(question)

        text = question.get("text", "")
        methodology = question.get("methodology", "")
        numeric_type = question.get("numeric_type", "")

        # Detect currency
        currency_match = re.search(r'[£$€¥₹]', text)
        has_currency_symbol = currency_match is not None
        currency = currency_match.group(0) if currency_match else "$"

        # Detect unit
        unit_match = re.search(r'per\s+([^,.?]+)', text, re.IGNORECASE)
        unit = unit_match.group(1).strip() if unit_match else ""

        # Check for Van Westendorp or explicit currency context
        is_van_westendorp = 'van_westendorp' in methodology.lower() or 'van westendorp' in text.lower()
        
        # Determine if this is a currency question
        is_currency_question = (
            is_van_westendorp or 
            numeric_type == 'currency' or
            (has_currency_symbol and any(kw in text.lower() for kw in ['price', 'cost', 'amount', 'budget', 'pay', 'spend']))
        )

        doc = self._get_document()
        if is_van_westendorp:
            paragraph = doc.add_paragraph()
            paragraph.add_run("Van Westendorp Price Sensitivity Question").bold = True

            guidance_paragraph = doc.add_paragraph()
            guidance_paragraph.add_run("Please think about the price point where you would have the described reaction to the product. Consider your local market conditions and personal budget.").italic = True

        # Render input based on question type
        input_paragraph = doc.add_paragraph()
        if is_currency_question:
            # Currency and amount input for currency questions
            currency_text = f"Currency: {currency}    Amount: {currency}_________"
            if unit:
                currency_text += f" per {unit}"
            input_paragraph.add_run(currency_text)
        else:
            # Generic number input for non-currency numeric questions
            numeric_text = "Value: _________"
            if unit:
                numeric_text += f" {unit}"
            input_paragraph.add_run(numeric_text)

        doc.add_paragraph("")

    def _render_likert(self, question: Dict[str, Any]) -> None:
        """Render Likert scale question."""
        self._add_question_header(question)
        self._add_required_indicator(question)

        # Add scale labels
        doc = self._get_document()
        scale_paragraph = doc.add_paragraph()
        scale_paragraph.add_run("Scale: ")

        scale_labels = ["Very Unlikely", "Unlikely", "Neutral", "Likely", "Very Likely"]
        for i, label in enumerate(scale_labels):
            if i > 0:
                scale_paragraph.add_run(" | ")
            scale_paragraph.add_run(f"{i+1}. {label}")

        # Add radio button options
        for i in range(1, 6):
            option_paragraph = doc.add_paragraph()
            option_paragraph.style = 'List Bullet'
            option_paragraph.add_run(f"○ {i}")

        doc.add_paragraph("")

    def _render_open_end(self, question: Dict[str, Any]) -> None:
        """Render open-ended text question."""
        self._add_question_header(question)
        self._add_required_indicator(question)

        # Add text input area
        doc = self._get_document()
        input_paragraph = doc.add_paragraph()
        input_paragraph.add_run("Response: ")

        # Add a line for writing
        response_paragraph = doc.add_paragraph()
        response_paragraph.add_run("_" * 50)

        doc.add_paragraph("")

    def _render_display_only(self, question: Dict[str, Any]) -> None:
        """Render display-only instruction block."""
        # Add instruction header
        doc = self._get_document()
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("📋 Programmer Instructions")
        run.bold = True
        run.font.size = Pt(12)

        # Add instruction content in a bordered paragraph
        instruction_paragraph = doc.add_paragraph()
        instruction_paragraph.style = 'Quote'
        instruction_paragraph.add_run(question.get("text", ""))

        # Add description if present
        if question.get("description"):
            desc_paragraph = doc.add_paragraph()
            desc_paragraph.add_run(f"Note: {question['description']}")
            desc_paragraph.style = 'Intense Quote'

        # Add metadata
        meta_paragraph = doc.add_paragraph()
        meta_paragraph.add_run("Type: Display Only | No response required")
        meta_paragraph.style = 'Caption'

        doc.add_paragraph("")

    def _render_single_open(self, question: Dict[str, Any]) -> None:
        """Render single-line open-ended question."""
        self._add_question_header(question)
        self._add_required_indicator(question)

        # Add single line input
        doc = self._get_document()
        input_paragraph = doc.add_paragraph()
        input_paragraph.add_run("Response: ")

        # Add a line for writing
        response_paragraph = doc.add_paragraph()
        response_paragraph.add_run("_" * 50)

        doc.add_paragraph("")

    def _render_multiple_open(self, question: Dict[str, Any]) -> None:
        """Render multi-line open-ended question."""
        self._add_question_header(question)
        self._add_required_indicator(question)

        # Add multi-line input area
        doc = self._get_document()
        input_paragraph = doc.add_paragraph()
        input_paragraph.add_run("Response: ")

        # Add multiple lines for writing
        for i in range(4):
            response_paragraph = doc.add_paragraph()
            response_paragraph.add_run("_" * 50)

        doc.add_paragraph("")

    def _render_open_ended(self, question: Dict[str, Any]) -> None:
        """Render open-ended question with larger text area."""
        self._add_question_header(question)
        self._add_required_indicator(question)

        # Add open-ended input area
        doc = self._get_document()
        input_paragraph = doc.add_paragraph()
        input_paragraph.add_run("Detailed Response: ")

        # Add multiple lines for detailed writing
        for i in range(5):
            response_paragraph = doc.add_paragraph()
            response_paragraph.add_run("_" * 50)

        doc.add_paragraph("")

    def _render_gabor_granger(self, question: Dict[str, Any]) -> None:
        """Render Gabor-Granger price sensitivity question."""
        self._add_question_header(question)
        self._add_required_indicator(question)

        # Extract price points from question text or options
        text = question.get("text", "")
        options = question.get("options", [])
        
        # Try to extract prices from text if options are not provided
        if not options:
            # Look for price patterns like $249, $299, etc.
            import re
            price_pattern = r'\$[\d,]+(?:\.\d{2})?'
            prices = re.findall(price_pattern, text)
            if prices:
                options = prices

        # Default price points if none found
        if not options:
            options = ["$249", "$299", "$349", "$399", "$449", "$499", "$549", "$599"]

        doc = self._get_document()
        
        # Add instruction text
        instruction_paragraph = doc.add_paragraph()
        instruction_paragraph.add_run("Please indicate your purchase likelihood at each price point:").italic = True
        
        # Create table with price points and likelihood scale
        table = doc.add_table(rows=len(options) + 1, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        table.cell(0, 0).text = "Price Point"
        table.cell(0, 1).text = "Purchase Likelihood"

        # Data rows with price points and radio buttons
        for i, price in enumerate(options):
            table.cell(i + 1, 0).text = price
            
            # Add likelihood scale options
            likelihood_cell = table.cell(i + 1, 1)
            likelihood_paragraph = likelihood_cell.paragraphs[0]
            likelihood_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add scale options (Definitely Not, Probably Not, Maybe, Probably Yes, Definitely Yes)
            scale_options = ["Definitely Not", "Probably Not", "Maybe", "Probably Yes", "Definitely Yes"]
            for j, scale_option in enumerate(scale_options):
                if j > 0:
                    likelihood_paragraph.add_run("  ")
                likelihood_paragraph.add_run(f"○ {scale_option}")
                if j < len(scale_options) - 1:
                    likelihood_paragraph.add_run("  |  ")

        doc.add_paragraph("")

    def _render_unsupported_question_type(self, question: Dict[str, Any]) -> None:
        """
        Render unsupported question types with a generic text area.
        This ensures DOCX generation doesn't fail for unknown question types.
        
        Args:
            question: Question data dictionary
        """
        self._add_question_header(question)
        self._add_required_indicator(question)

        # Add warning note
        doc = self._get_document()
        warning_paragraph = doc.add_paragraph()
        warning_run = warning_paragraph.add_run(f"⚠️ Unsupported question type: {question.get('type', 'unknown')}")
        warning_run.font.size = Pt(9)
        warning_run.font.color.rgb = None  # Use default color

        # Add generic text input area
        input_paragraph = doc.add_paragraph()
        input_paragraph.add_run("Response: ")

        # Add lines for writing
        for i in range(3):
            response_paragraph = doc.add_paragraph()
            response_paragraph.add_run("_" * 50)

        doc.add_paragraph("")


# Register the DOCX renderer
export_registry.register_renderer("docx", DocxSurveyRenderer)